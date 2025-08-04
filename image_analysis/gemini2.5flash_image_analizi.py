import base64
import google.generativeai as genai
from PIL import Image
import io
from docx import Document
from datetime import datetime

# Gemini API anahtarı (kendi API anahtarınızı buraya ekleyin)
GEMINI_API_KEY = "AIzaSyAyjFX1qokUL0EJ41LTjZZS7jJmq87cxjM"  # Google Cloud'dan aldığınız API anahtarını buraya yazın
genai.configure(api_key=GEMINI_API_KEY)

# Gemini 2.5 Flash modelini yükle
model = genai.GenerativeModel("gemini-2.5-flash")

# Görseli base64 formatına çevirme fonksiyonu
def image_to_base64(image_path):
    try:
        with Image.open(image_path) as img:
            # Görseli JPEG formatına çevir
            buffered = io.BytesIO()
            img.save(buffered, format="JPEG")
            # Base64'e kodla
            return base64.b64encode(buffered.getvalue()).decode("utf-8")
    except FileNotFoundError:
        print(f"Hata: Görsel dosyası bulunamadı: {image_path}")
        return None

# Görsel dosya yolu (lütfen burayı kendi görsel yolunuzla değiştirin)
image_path = "1747162292672.jpg"  # Örnek: "candidate_interview.jpg"
image_base64 = image_to_base64(image_path)

if image_base64 is None:
    print("Görsel yüklenemedi. Lütfen doğru dosya yolunu belirtin.")
    exit()

# Prompt
prompt = """
Aşağıdaki görsel, bir adayın mülakat anında çekilmiş tek karelik bir görüntüsüdür.
Görsele bakarak sadece gözlemlenebilir unsurlar üzerinden aşağıdaki analizleri yap:

Kıyafet Analizi:
– Kıyafetin profesyonellik düzeyi, temizliği, uyumu, uygunluğu

Duruş/Postür:
– Dik duruş, kamburluk, omuz hizası, oturuş şekli

Beden Dili:
– Eller nerede? Açık mı, kapalı mı? Gerginlik, rahatlık belirtisi var mı?

Yüz İfadesi:
– Nötr, gergin, gülümseyen, stresli, odaklı, rahat gibi ifadeler

Genel İlk İzlenim:
– Görselden yansıyan özgüven, hazırlık düzeyi, profesyonellik

Görselde ses, konuşma, davranış bilgisi yoktur. Bu yüzden yalnızca görüntüye dayalı olarak:
Kişi hakkında oluşturulabilecek görsel temelli bir ilk izlenim raporu hazırla.
Her yorumun görselde neye dayandığını kısa şekilde belirt. Her analiz için açıklamalar 1-2 cümleyle sınırlı olmalıdır.
"""

# Görseli ve prompt'u Gemini API'sine gönder
try:
    response = model.generate_content([
        prompt,
        {
            "mime_type": "image/jpeg",
            "data": image_base64
        }
    ])
    analysis = response.text
except Exception as e:
    analysis = f"Hata: Gemini API'sine bağlanılamadı. Hata mesajı: {e}\nLütfen API anahtarının doğru olduğundan ve Gemini 2.5 Flash modeline erişiminizin olduğundan emin olun."

# Genel yorum (simüle edilmiş, görsele bağlı olarak güncellenecek)
general_comment = """
Adayın kıyafeti, duruşu ve beden dili, profesyonel ve özgüvenli bir izlenim sunuyor. Hafif gülümseme ve rahat postür, mülakata iyi hazırlanmış olduğunu gösteriyor.
"""

# Word dosyası oluştur
doc = Document()

# Başlık
doc.add_heading("Görsel Temelli İlk İzlenim Raporu", 0)

# Tarih ve saat
current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
doc.add_paragraph(f"Tarih: {current_time}")

# Analiz sonuçları bölümü
doc.add_heading("Analiz Sonuçları", level=1)
doc.add_paragraph(analysis, style="Normal")

# Genel yorum bölümü
doc.add_heading("Genel Yorum", level=1)
doc.add_paragraph(general_comment, style="Normal")

# Word dosyasını kaydet
output_file = "image_analysis_report.docx"
doc.save(output_file)
print(f"Analiz raporu '{output_file}' olarak kaydedildi.")