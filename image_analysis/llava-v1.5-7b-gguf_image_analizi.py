import base64
import requests
from PIL import Image
import io
from docx import Document
from datetime import datetime
import os

# --- AYARLAR ---

# LM Studio yerel API adresi
API_URL = "http://localhost:1234/v1/chat/completions"

# Analiz edilecek görselin yolu (Lütfen bu yolu kendi dosyanızla değiştirin)
# Örnek: "C:/Users/Kullanici/Desktop/aday_fotografi.jpg"
IMAGE_PATH = "1747162292672.jpg"

# LM Studio'da yüklü olan GGUF modelinin tam adı
MODEL_NAME = "llava-v1.5-7b-gguf"

# Çıktı Word dosyasının adı
OUTPUT_FILENAME = "Tamamen_Turkce_Mülakat_Analizi.docx"


# --- FONKSİYONLAR ---

def image_to_base64(image_path):
    """Verilen yoldaki bir görseli base64 formatına çevirir."""
    if not os.path.exists(image_path):
        print(f"Hata: Görsel dosyası bulunamadı: {image_path}")
        return None
    try:
        with Image.open(image_path) as img:
            buffered = io.BytesIO()
            # RGBA gibi formatları standart RGB'ye çevirerek uyumluluğu artır
            if img.mode in ('RGBA', 'P'):
                img = img.convert('RGB')
            img.save(buffered, format="JPEG")
            return base64.b64encode(buffered.getvalue()).decode("utf-8")
    except Exception as e:
        print(f"Hata: Görsel işlenirken bir sorun oluştu: {e}")
        return None


def create_analysis_report(analysis_content, image_path):
    """Verilen analiz içeriğiyle bir Word raporu oluşturur."""
    try:
        doc = Document()
        doc.add_heading("Görsel Temelli İlk İzlenim Raporu", 0)

        current_time = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        doc.add_paragraph(f"Rapor Tarihi: {current_time}")
        doc.add_paragraph(f"Analiz Edilen Görsel: {os.path.basename(image_path)}\n")

        doc.add_heading("Yapay Zeka Analiz Sonuçları", level=1)
        doc.add_paragraph(analysis_content, style="Normal")

        doc.save(OUTPUT_FILENAME)
        print(f"\n✓ Analiz raporu başarıyla '{OUTPUT_FILENAME}' olarak kaydedildi.")
    except Exception as e:
        print(f"Hata: Word dosyası oluşturulurken bir sorun yaşandı: {e}")


# --- ANA KOD BLOKU ---

if __name__ == "__main__":
    # 1. Görseli Base64 formatına çevir
    image_base64_str = image_to_base64(IMAGE_PATH)

    if not image_base64_str:
        exit()

    # 2. LLaVA için TAMAMEN TÜRKÇE çıktı üretecek şekilde optimize edilmiş prompt
    prompt_text = """
Sen uzman bir İnsan Kaynakları analistisin. Görevin, bir iş görüşmesindeki adayın bu tek kare fotoğrafını analiz etmek.
Tüm çıktıyı baştan sona Türkçe olarak yazmalısın.
Sadece fotoğrafta gördüklerine dayanarak, aşağıdaki başlıklar altında yapılandırılmış bir analiz sun. Her başlık için gördüklerini anlat ve kısa bir yorum ekle.

* **Kıyafet Analizi:** Kıyafetin profesyonelliğini, temizliğini ve vücuda uygunluğunu anlat.
* **Duruş (Postür) Analizi:** Adayın oturuş pozisyonunu, sırtının ve omuzlarının duruşunu tarif et.
* **Beden Dili:** Ellerin pozisyonunu açıkla. Gözle görülür bir gerginlik veya rahatlık belirtisi olup olmadığını not et.
* **Yüz İfadesi:** Genel yüz ifadesini (örneğin; nötr, odaklanmış, gülümseyen, gergin) tanımla.
* **Genel İlk İzlenim:** Tüm görsel kanıtlara dayanarak, adayın özgüveni, hazırlık durumu ve profesyonelliği hakkındaki ilk izlenimini özetle.

Analizinin sonunda, tüm bu gözlemlerini birleştiren bütüncül bir "Genel Değerlendirme" paragrafını Türkçe olarak yaz.
"""

    # 3. LM Studio API'sine gönderilecek payload'ı oluştur
    payload = {
        "model": MODEL_NAME,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt_text},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64_str}"}}
                ]
            }
        ],
        "max_tokens": 2048,
        "temperature": 0.4
    }

    # 4. API isteğini gönder ve analizi al
    print("LM Studio API'sine bağlanılıyor ve Türkçe analiz isteniyor...")
    try:
        response = requests.post(API_URL, json=payload, timeout=120)
        response.raise_for_status()

        result = response.json()
        analysis_result = result["choices"][0]["message"]["content"]

        print("✓ Analiz başarıyla alındı.")

        # 5. Word raporunu oluştur
        create_analysis_report(analysis_result, IMAGE_PATH)

    except requests.exceptions.RequestException as e:
        print(f"\nHata: LM Studio API'sine bağlanılamadı. Lütfen kontrol edin: {e}")
        print("İpuçları: LM Studio'nun çalıştığından ve 'Server' sekmesinden sunucuyu başlattığınızdan emin olun.")
    except (KeyError, IndexError):
        print(f"\nHata: API'den gelen yanıt beklenmedik bir formatta. Yanıt: {response.text}")
    except Exception as e:
        print(f"Beklenmedik bir hata oluştu: {e}")