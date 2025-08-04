import base64
import requests
import asyncio
from PIL import Image
import io
from docx import Document
from datetime import datetime
from googletrans import Translator, LANGUAGES

# LM Studio API URL'si
API_URL = "http://localhost:1234/v1/chat/completions"

# Analiz edilecek görselin yolu (Lütfen bu yolu kendi dosyanızla değiştirin)
IMAGE_PATH = "1747162292672.jpg"  # Örnek: "C:/Users/Kullanici/Desktop/aday_fotografi.jpg"

# LM Studio'da yüklü olan GGUF modelinin tam adı
MODEL_NAME = "BakLLaVA1-MistralLLaVA-7B-Q5_K_M"  # LM Studio'daki model adıyla eşleşmeli

# Görseli base64 formatına çevirme fonksiyonu
def image_to_base64(image_path):
    try:
        with Image.open(image_path) as img:
            # Görseli JPEG formatına çevir
            buffered = io.BytesIO()
            img.save(buffered, format="JPEG")
            # Base64'e kodla
            return base64.b64encode(buffered.getvalue()).decode("utf-8")
    except FileNotFoundError:
        print(f"Hata: Görsel dosyası bulunamadı: {image_path}")
        return None

# Görseli base64'e çevir
image_base64 = image_to_base64(IMAGE_PATH)

if image_base64 is None:
    print("Görsel yüklenemedi. Lütfen doğru dosya yolunu belirtin.")
    exit()

# Prompt Şablonu (İngilizce)
prompt_template = """
The following image is a single photograph of a candidate during a job interview. Act as an HR specialist and analyze the image based on observable elements including clothing suitability for an interview, background lighting quality, hair and beard condition, facial expression, posture, and any other relevant visual cues. Provide all responses in English.
Based on these observations, write a single paragraph evaluating the candidate's overall suitability, addressing whether the clothing is appropriate for an interview, the background lighting is adequate, and the hair and beard are well-groomed (e.g., clean-shaven, neat hair, or otherwise). Include comments on professionalism, confidence, and preparedness inferred from facial expression and posture.
There is no sound, speech, or behavioral information in the image. Therefore, base the evaluation solely on the visual content of the photograph.
Keep the comment to one paragraph and indicate what each observation is based on in the image.
"""

# LM Studio API'sine istek gönderme
headers = {
    "Content-Type": "application/json"
}
payload = {
    "model": MODEL_NAME,
    "messages": [
        {
            "role": "user",
            "content": [
                {"type": "text", "text": prompt_template},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}}
            ]
        }
    ],
    "max_tokens": 300
}

try:
    response = requests.post(API_URL, json=payload, headers=headers, timeout=30)
    response.raise_for_status()
    analysis = response.json()["choices"][0]["message"]["content"]
except requests.exceptions.RequestException as e:
    analysis = f"Hata: LM Studio API'sine bağlanılamadı. Hata mesajı: {e}\nLütfen LM Studio'nun çalıştığını ve doğru portta (1234) olduğundan emin olun."

# Asenkron çeviri fonksiyonu
async def translate_text(text, src='en', dest='tr'):
    translator = Translator()
    translation = await translator.translate(text, src=src, dest=dest)
    return translation.text

# Çeviriyi senkronize etmek için event loop kullan
loop = asyncio.get_event_loop()
translated_analysis = loop.run_until_complete(translate_text(analysis))

# Word dosyası oluştur
doc = Document()

# Başlık
doc.add_heading("Görsel Temelli Mülakat Analiz Raporu", 0)

# Tarih ve saat
current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
doc.add_paragraph(f"Tarih: {current_time}")

# Analiz sonuçları bölümü
doc.add_heading("Değerlendirme", level=1)
doc.add_paragraph(translated_analysis, style="Normal")

# Word dosyasını kaydet
output_file = "mülakat_analiz_raporu.docx"
doc.save(output_file)
print(f"Analiz raporu '{output_file}' olarak kaydedildi.")