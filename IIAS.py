# -*- coding: utf-8 -*-

# --- GEREKLİ KÜTÜPHANELER ---
import os
import json
import requests
import spacy
import cv2
import pytesseract
from moviepy import VideoFileClip
from docx import Document
from elevenlabs import ElevenLabs
from langchain.text_splitter import RecursiveCharacterTextSplitter
from deepface import DeepFace
from collections import Counter
import warnings
import numpy as np
from PIL import Image
import base64
import google.generativeai as genai
import io
from datetime import datetime

# DeepFace ve TensorFlow uyarılarını gizle
warnings.filterwarnings("ignore", category=UserWarning)

# pytesseract yolunu belirt (Windows için, sisteminize göre güncelleyin)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Windows için
# Linux/Mac için bu satırı kaldırın veya yorum satırı yapın

# --- KONFİGÜRASYON ---
ELEVENLABS_API_KEY = "sk_6bd978abf5bfe173b11565c80fb8d0233d0cd6015b71c4c3"
GEMINI_API_KEY = "AIzaSyAyjFX1qokUL0EJ41LTjZZS7jJmq87cxjM"  # Google Cloud'dan aldığınız API anahtarını buraya yazın
INPUT_VIDEO_FILE = "video1723838072.mp4"
TEMP_AUDIO_FILE = "gecici_ses.wav"
TRANSCRIPT_DOCX_FILE = "mulakat_transkripti.docx"
SANIYEDE_ANALIZ_SAYISI = 2
TEMP_FRAME_FILE = "temp_frame.jpg"
TEMP_FACE_FILE = "temp_face.jpg"
LM_STUDIO_API_URL = "http://localhost:1234/v1/chat/completions"
MODEL_NAME = "mistral-7b-instruct-v0.3.Q4_0"
FINAL_ANALYSIS_DOCX_FILE = "analiz_sonucu.docx"
DURATION_THRESHOLD = 20

# DeepFace duygu etiketlerini Türkçeye çevirme sözlüğü
DUYGU_SOZLUGU = {
    'angry': 'ÖFKELİ',
    'disgust': 'TİKSİNMİŞ',
    'fear': 'KORKMUŞ',
    'happy': 'MUTLU',
    'sad': 'ÜZGÜN',
    'surprise': 'ŞAŞKIN',
    'neutral': 'DOĞAL'
}

# --- YENİ PROMPT'LAR ---
PROMPT_SCORING_DETAILS = """
Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.

BÖLÜM 1: PUANLAMA TABLOSU
• Her kriteri, madde imi (•) ile başlayan ayrı bir satırda ve şu formatta yaz:
• Kriter Adı: (Puan/5) - {candidate_name}'in [puanın nedenini açıklayan kısa ve tanımlayıcı bir cümle].
• Değerlendirilecek Kriterler:
• İletişim Becerisi
• Motivasyon ve Tutku
• Kültürel Uyum
• Analitik/Düşünsel Beceriler
• Profesyonel Tutum
• Geçmiş Deneyim Uyumu
• Liderlik ve Girişimcilik
• Zayıflıklarla Başa Çıkma Yetisi
• Uzun Vadeli Potansiyel
• Genel Etki / İzlenim
• Analiz Sonu:
• Genel Ortalama Puan: Tüm puanların ortalamasını, ondalık ayraç olarak virgül kullanarak hesapla. Örnek: (3,86/5)
• İK Genel Yorum: Adayın genel potansiyelini ve ana bulguları özetleyen birkaç cümlelik bir paragraf yaz."""

PROMPT_RECRUITER_DETAILS = """
Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.

BÖLÜM 2: RECRUITER NOTU TALİMATLARI

• Analizini, aşağıdaki altı başlığın tamamını madde imi (•) ile başlayan ayrı satırlar olarak yapılandır.
• Her başlığın altına, mülakat özetinden çıkardığın somut bilgilere (projeler, deneyimler, yetenekler) dayanarak detaylı ve profesyonel bir metin yaz. Adayın adını metin içinde uygun yerlerde kullan.
• Başlıklar ve İçerikleri:
• Aday Adı: Mülakat metninden adayın adını çıkar.
• Pozisyon: Mülakat metninden adayın başvurduğu pozisyonu belirle.
• Genel Yorum: Adayın geçmişi, deneyim süresi ve genel performansı hakkında özet bir paragraf yaz.
• Dikkat Çeken Güçlü Yönler: Adayın öne çıkan teknik veya sosyal yeteneklerini, projelerden örnekler vererek anlat.
• Geliştirme Alanları: Adayın hangi konularda kendini geliştirebileceğini ve potansiyel gelişim alanlarını belirt.
• Değerlendirme Önerisi: Aday için bir sonraki adımları (ikinci görüşme, teknik test vb.) ve gelişimini destekleyecek önerileri (eğitim, kurs vb.) içeren bir paragraf yaz.
"""

# ==============================================================================
# --- YARDIMCI FONKSİYONLAR ---
# ==============================================================================

def extract_candidate_name_from_text(transcript_file_path):
    """
    Transkript dosyasından adayın ismini LLM ile çıkarır.
    """
    try:
        interview_text = read_text_from_docx(transcript_file_path)
        if not interview_text or not interview_text.strip():
            print(f"HATA: '{transcript_file_path}' dosyası bulunamadı veya boş.")
            return None

        prompt_name = f"""
        Aşağıdaki mülakat metninden adayın ismini çıkar. Sadece ismi döndür, başka bir şey yazma.
        Örneğin: "Merhaba, ben Can Bey" -> Can Bey

        --- MÜLAKAT METNİ ---
        {interview_text}
        --- İSİM ---
        """
        candidate_name = get_llm_analysis(prompt_name, MODEL_NAME)
        if candidate_name and candidate_name.strip():
            print(f"Metinden tespit edilen isim: {candidate_name}")
            return candidate_name.strip()
        else:
            print("Uyarı: Metinden isim tespit edilemedi.")
            return None
    except Exception as e:
        print(f"HATA: Metinden isim çıkarılırken hata: {e}")
        return None

def extract_frame_and_name(video_path, text_name, max_duration=120.0):
    """
    Videonun ilk 120 saniyesinde her 2 saniyede bir sol alt ve sağ alt köşelerde OCR ile isim çıkarır,
    metin ismiyle eşleşen ilk karede yüzü kırpar.
    """
    try:
        print(f"Videonun ilk {max_duration} saniyesinde her 2 saniyede bir isim kontrol ediliyor...")
        video = cv2.VideoCapture(video_path)
        if not video.isOpened():
            print(f"HATA: '{video_path}' videosu açılamadı.")
            return None, None
        fps = video.get(cv2.CAP_PROP_FPS) or 30  # FPS alınamazsa varsayılan 30
        max_frames = int(max_duration * fps)
        frame_interval = int(fps * 2)  # Her 2 saniyede bir kare
        frame_num = 0
        final_name = None
        face_image_path = None

        # Video boyutlarını al
        frame_width = int(video.get(cv2.CAP_PROP_FRAME_WIDTH))
        frame_height = int(video.get(cv2.CAP_PROP_FRAME_HEIGHT))
        region_height = int(frame_height * 0.30)  # Alt %30
        region_width = int(frame_width * 0.30)   # Sol/sağ %30
        left_bottom = (0, frame_height - region_height, region_width, frame_height)
        right_bottom = (frame_width - region_width, frame_height - region_height, frame_width, frame_height)

        while frame_num < max_frames:
            ret, frame = video.read()
            if not ret:
                print("HATA: Video karesi alınamadı.")
                break

            if frame_num % frame_interval == 0:
                # Sol alt ve sağ alt bölgeleri kırp
                left_crop = frame[left_bottom[1]:left_bottom[3], left_bottom[0]:left_bottom[2]]
                right_crop = frame[right_bottom[1]:right_bottom[3], right_bottom[0]:right_bottom[2]]
                cv2.imwrite(TEMP_FRAME_FILE + "_left.jpg", left_crop)
                cv2.imwrite(TEMP_FRAME_FILE + "_right.jpg", right_crop)
                print(f"Kareler kaydedildi: sol alt ve sağ alt (zaman: {frame_num/fps:.2f} saniye).")

                # OCR ile isim çıkar
                ocr_names = []
                for region_file in [TEMP_FRAME_FILE + "_left.jpg", TEMP_FRAME_FILE + "_right.jpg"]:
                    image = Image.open(region_file)
                    ocr_text = pytesseract.image_to_string(image, lang='tur', config='--psm 7')
                    for line in ocr_text.split('\n'):
                        line = ''.join(c for c in line if c.isalnum() or c.isspace()).strip()
                        if line and len(line.split()) <= 3:
                            ocr_names.append(line)
                print(f"OCR ile tespit edilen isimler: {ocr_names if ocr_names else 'Tespit edilemedi'}")

                # Metin ismiyle eşleştir
                text_name_lower = text_name.lower().strip()
                text_first_name = text_name_lower.split()[0]
                for ocr_name in ocr_names:
                    if text_first_name in ocr_name.lower().strip():
                        final_name = text_name
                        print(f"İsim eşleşti: {text_name}")
                        # Yüz tespiti
                        try:
                            analysis = DeepFace.analyze(frame, actions=['emotion'], enforce_detection=False, silent=True)
                            if isinstance(analysis, list) and len(analysis) > 0:
                                face_region = analysis[0]['region']
                                x, y, w, h = face_region['x'], face_region['y'], face_region['w'], face_region['h']
                                face = frame[y:y+h, x:x+w]
                                cv2.imwrite(TEMP_FACE_FILE, face)
                                face_image_path = TEMP_FACE_FILE
                                print(f"Adayın yüzü '{TEMP_FACE_FILE}' olarak kaydedildi.")
                                break
                        except Exception:
                            print("Uyarı: Yüz tespit edilemedi.")
                if final_name:
                    break  # Eşleşme bulundu, döngüden çık

            frame_num += 1

        video.release()
        for region_file in [TEMP_FRAME_FILE + "_left.jpg", TEMP_FRAME_FILE + "_right.jpg"]:
            if os.path.exists(region_file):
                os.remove(region_file)
                print(f"Geçici dosya '{region_file}' silindi.")

        if not final_name or not face_image_path:
            print("Uyarı: İlk 120 saniyede eşleşen isim veya yüz tespit edilemedi.")
        return final_name, face_image_path
    except Exception as e:
        print(f"HATA: Kare çıkarma veya analiz sırasında hata: {e}")
        return None, None

def compare_names(text_name, ocr_names):
    """
    Metinden alınan isimle OCR'den alınan isimleri karşılaştırır, eşleşen ismi döndürür.
    """
    if not text_name or not ocr_names:
        return text_name or "Aday"

    text_name_lower = text_name.lower().strip()
    for ocr_name in ocr_names:
        if ocr_name.lower().strip() == text_name_lower:
            print(f"İsim eşleşti: {text_name}")
            return text_name
    print(f"Uyarı: OCR isimleri ({ocr_names}) metin ismi ({text_name}) ile eşleşmedi.")
    return None

def analyze_character_from_image(face_image_path):
    """
    Gemini 2.5 Flash ile adayın yüz görüntüsüne dayalı karakter analizi yapar.
    """
    try:
        print(f"Adayın yüz görüntüsü '{face_image_path}' Gemini 2.5 Flash ile analiz ediliyor...")

        # Gemini API'sini yapılandır
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel("gemini-2.5-flash")

        # Görseli base64 formatına çevir
        def image_to_base64(image_path):
            try:
                with Image.open(image_path) as img:
                    buffered = io.BytesIO()
                    img.save(buffered, format="JPEG")
                    return base64.b64encode(buffered.getvalue()).decode("utf-8")
            except FileNotFoundError:
                print(f"Hata: Görsel dosyası bulunamadı: {image_path}")
                return None

        image_base64 = image_to_base64(face_image_path)
        if image_base64 is None:
            print("Görsel yüklenemedi. Analiz atlanıyor.")
            return None

        # Gemini için prompt
        prompt = """
        Aşağıdaki görsel, bir adayın mülakat anında çekilmiş tek karelik bir görüntüsüdür.
        Görsele bakarak sadece gözlemlenebilir unsurlar üzerinden aşağıdaki analizleri yap:

        Kıyafet Analizi:
        – Kıyafetin profesyonellik düzeyi, temizliği, uyumu, uygunluğu

        Duruş/Postür:
        – Dik duruş, kamburluk, omuz hizası, oturuş şekli

        Beden Dili:
        – Eller nerede? Açık mı, kapalı mı? Gerginlik, rahatlık belirtisi var mı?

        Yüz İfadesi:
        – Nötr, gergin, gülümseyen, stresli, odaklı, rahat gibi ifadeler

        Genel İlk İzlenim:
        – Görselden yansıyan özgüven, hazırlık düzeyi, profesyonellik

        Görselde ses, konuşma, davranış bilgisi yoktur. Bu yüzden yalnızca görüntüye dayalı olarak:
        Kişi hakkında oluşturulabilecek görsel temelli bir ilk izlenim raporu hazırla.
        Her yorumun görselde neye dayandığını kısa şekilde belirt. Her analiz için açıklamalar 1-2 cümleyle sınırlı olmalıdır.
        """

        # Görseli ve prompt'u Gemini API'sine gönder
        response = model.generate_content([
            prompt,
            {
                "mime_type": "image/jpeg",
                "data": image_base64
            }
        ])
        analysis = response.text
        print("Gemini 2.5 Flash görüntü analizi başarıyla tamamlandı.")
        return analysis
    except Exception as e:
        print(f"HATA: Görüntü analizi sırasında hata: {e}")
        return None

def videodaki_duygulari_analiz_et(video_path, saniyede_kontrol=2):
    """Videoyu analiz eder ve zaman damgalarıyla birlikte dominant duyguları listeler."""
    print("\nVideo üzerinden duygu analizi başlatılıyor (bu işlem videonun uzunluğuna göre sürebilir)...")
    video = cv2.VideoCapture(video_path)
    fps = video.get(cv2.CAP_PROP_FPS)
    frame_interval = int(fps / saniyede_kontrol) if saniyede_kontrol > 0 and fps > 0 else int(fps)

    duygu_zaman_cizelgesi = []
    frame_num = 0

    while video.isOpened():
        ret, frame = video.read()
        if not ret:
            break

        if frame_num % frame_interval == 0:
            try:
                analysis = DeepFace.analyze(
                    frame,
                    actions=['emotion'],
                    enforce_detection=False,
                    silent=True
                )
                if isinstance(analysis, list) and len(analysis) > 0:
                    dominant_emotion_en = analysis[0]['dominant_emotion']
                    dominant_emotion_tr = DUYGU_SOZLUGU.get(dominant_emotion_en, dominant_emotion_en.upper())
                    timestamp = frame_num / fps
                    duygu_zaman_cizelgesi.append({'zaman': timestamp, 'duygu': dominant_emotion_tr})
            except Exception:
                pass
        frame_num += 1

    video.release()
    print("Duygu analizi tamamlandı.")
    return duygu_zaman_cizelgesi

def konusmacilari_ayir_ve_cumlele(words_data, duygu_cizelgesi):
    """
    ElevenLabs'ten gelen kelime bazlı diarizasyon verisini işler.
    Konuşmacı etiketlerine göre metni gruplar, spaCy ile cümlelere ayırır ve duygu analizi ekler.
    """
    if not words_data:
        print("\nKonuşmacı ayırmak için kelime verisi bulunamadı.")
        return []

    speaker_map = {}
    next_speaker_num = 1
    current_speaker_id_raw = None
    current_speaker_mapped_label = None
    current_speaker_words_buffer = []
    all_speaker_utterances_combined = []

    for word_obj in words_data:
        word_text = word_obj.text
        raw_speaker_id = word_obj.speaker_id
        start_time = word_obj.start
        end_time = word_obj.end

        if raw_speaker_id not in speaker_map:
            speaker_map[raw_speaker_id] = f"Konuşmacı {next_speaker_num}"
            next_speaker_num += 1

        mapped_speaker_label = speaker_map[raw_speaker_id]

        if current_speaker_id_raw is None:
            current_speaker_id_raw = raw_speaker_id
            current_speaker_mapped_label = mapped_speaker_label
        elif raw_speaker_id != current_speaker_id_raw:
            if current_speaker_words_buffer:
                all_speaker_utterances_combined.append({
                    'speaker': current_speaker_mapped_label,
                    'text_combined': " ".join([w['text'] for w in current_speaker_words_buffer]).strip(),
                    'start_time': current_speaker_words_buffer[0]['start'],
                    'end_time': current_speaker_words_buffer[-1]['end']
                })
            current_speaker_id_raw = raw_speaker_id
            current_speaker_mapped_label = mapped_speaker_label
            current_speaker_words_buffer = []

        current_speaker_words_buffer.append({'text': word_text, 'start': start_time, 'end': end_time})

    if current_speaker_words_buffer:
        all_speaker_utterances_combined.append({
            'speaker': current_speaker_mapped_label,
            'text_combined': " ".join([w['text'] for w in current_speaker_words_buffer]).strip(),
            'start_time': current_speaker_words_buffer[0]['start'],
            'end_time': current_speaker_words_buffer[-1]['end']
        })

    final_diarized_sentences = []
    try:
        nlp = spacy.load("xx_ent_wiki_sm")
        if "sentencizer" not in nlp.pipe_names:
            nlp.add_pipe("sentencizer")
    except OSError:
        print("\nspaCy modeli 'xx_ent_wiki_sm' bulunamadı.")
        print("Lütfen 'python -m spacy download xx_ent_wiki_sm' komutunu çalıştırın.")
        return []
    except Exception as e:
        print(f"spaCy modeli yüklenirken bir hata oluştu: {e}")
        return []

    for entry in all_speaker_utterances_combined:
        doc = nlp(entry['text_combined'])
        sentences_for_speaker = [sent.text.strip() for sent in doc.sents if sent.text.strip()]
        start_time = entry['start_time']
        end_time = entry['end_time']
        duration = end_time - start_time
        sentences_count = len(sentences_for_speaker)

        if sentences_count == 0:
            continue

        for i, sent in enumerate(sentences_for_speaker):
            sent_start = start_time + (i * duration / sentences_count)
            sent_end = start_time + ((i + 1) * duration / sentences_count)
            ilgili_duygular = [d['duygu'] for d in duygu_cizelgesi if sent_start <= d['zaman'] <= sent_end]
            dominant_duygu = Counter(ilgili_duygular).most_common(1)[0][0] if ilgili_duygular else "BELİRSİZ"

            final_diarized_sentences.append({
                "konusmaci": entry['speaker'],
                "diyalog": sent,
                "duygu": dominant_duygu,
                "baslangic": sent_start,
                "bitis": sent_end
            })

    return final_diarized_sentences

def read_text_from_docx(file_path):
    """
    Bir .docx dosyasındaki tüm metni okur ve tek bir metin bloğu olarak döndürür.
    """
    try:
        doc = Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        return '\n'.join(full_text)
    except Exception as e:
        print(f"HATA: '{file_path}' dosyası okunurken hata oluştu: {e}")
        return None

def get_llm_analysis(prompt, model_name):
    """
    LM Studio API aracılığıyla yerel LLM'e bir prompt gönderir ve analizi alır.
    """
    payload = {
        "model": model_name,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7,
        "max_tokens": 2048
    }
    try:
        response = requests.post(LM_STUDIO_API_URL, json=payload)
        response.raise_for_status()
        response_data = response.json()
        if "choices" in response_data and response_data["choices"]:
            return response_data["choices"][0]["message"]["content"]
        else:
            print(f"HATA: API yanıtında 'choices' anahtarı bulunamadı veya boş. Yanıt: {response_data}")
            return None
    except requests.exceptions.ConnectionError:
        print("HATA: LM Studio API sunucusuna bağlanılamadı. Lütfen sunucunun çalıştığından emin olun.")
        return None
    except requests.exceptions.HTTPError as http_err:
        print(f"HATA: API isteği sırasında HTTP hatası: {http_err}\nYanıt İçeriği: {response.text}")
        return None
    except Exception as e:
        print(f"HATA: Analiz sırasında beklenmedik bir hata oluştu: {e}")
        return None

def write_analysis_to_docx(file_path, analysis_scoring, analysis_recruiter, image_analysis, candidate_name):
    """
    Analiz sonuçlarını yeni bir .docx dosyasına yazar, yalnızca adayın görüntü analizi dahil.
    """
    try:
        doc = Document()
        doc.add_heading('Mülakat Analizi Sonucu', level=0)
        doc.add_heading('1. Aday Değerlendirme Puanlama Tablosu', level=1)
        for line in analysis_scoring.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_paragraph()
        doc.add_heading('2. Recruiter Notu', level=1)
        for line in analysis_recruiter.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_paragraph()
        doc.add_heading('3. Görüntü Tabanlı Karakter Analizi', level=1)
        if image_analysis:
            for line in image_analysis.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph(f"{candidate_name} için görüntü tabanlı analiz yapılamadı.")
        doc.save(file_path)
        print(f"\nAnaliz sonuçları '{file_path}' dosyasına başarıyla kaydedildi.")
    except Exception as e:
        print(f"HATA: Sonuçlar dosyaya yazılırken hata oluştu: {e}")

def adim_2_metin_analizi_yap(transcript_file_path, candidate_name="Aday"):
    """
    ADIM 2: Transkripti DOCX dosyasından okur, LLM'e analiz için gönderir ve sonuçları döndürür.
    """
    print(f"--- ADIM 2: '{transcript_file_path}' Dosyasındaki Metin Analiz Ediliyor ---")

    interview_text = read_text_from_docx(transcript_file_path)
    if not interview_text or not interview_text.strip():
        print(f"HATA: '{transcript_file_path}' dosyası bulunamadı, bozuk veya boş. Analiz yapılamıyor.")
        return None, None

    print("Mülakat transkripti başarıyla okundu.")

    print(f"'{MODEL_NAME}' modeline 'Aday Değerlendirme Puanlama Tablosu' için istek gönderiliyor...")
    prompt_scoring = PROMPT_SCORING_DETAILS.format(candidate_name=candidate_name)
    prompt_scoring = f"""
    Aşağıdaki mülakat metnini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et. 
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:
    {prompt_scoring}

    --- MÜLAKAT METNİ ---
    {interview_text}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_scoring = get_llm_analysis(prompt_scoring, MODEL_NAME)
    if not analysis_scoring:
        print("Puanlama analizi alınamadı. İşlem durduruluyor.")
        return None, None
    print("Puanlama analizi başarıyla tamamlandı.")

    print(f"'{MODEL_NAME}' modeline 'Recruiter Notu' için istek gönderiliyor...")
    prompt_recruiter = PROMPT_RECRUITER_DETAILS.format(candidate_name=candidate_name)
    prompt_recruiter = f"""
    Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:
    {prompt_recruiter}

    --- MÜLAKAT METNİ ---
    {interview_text}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_recruiter = get_llm_analysis(prompt_recruiter, MODEL_NAME)
    if not analysis_recruiter:
        print("Recruiter notu analizi alınamadı. İşlem durduruluyor.")
        return None, None
    print("Recruiter notu analizi başarıyla tamamlandı.")

    return analysis_scoring, analysis_recruiter

def adim_2_metin_analizi_chunk(transcript_file_path, candidate_name="Aday"):
    """
    ADIM 2 (Chunk Tabanlı): Transkripti DOCX dosyasından okur, metni parçalara böler,
    her parça için özet oluşturur, özetleri birleştirir ve tek bir nihai analiz yapar.
    """
    print(f"--- ADIM 2 (Chunk Tabanlı): '{transcript_file_path}' Dosyasındaki Metin Analiz Ediliyor ---")

    interview_text = read_text_from_docx(transcript_file_path)
    if not interview_text or not interview_text.strip():
        print(f"HATA: '{transcript_file_path}' dosyası bulunamadı, bozuk veya boş. Analiz yapılamıyor.")
        return None, None

    print("Mülakat transkripti başarıyla okundu.")

    print("Metin parçalara ayrılıyor...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=4000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(interview_text)
    print(f"Metin {len(chunks)} parçaya bölündü.")

    chunk_summaries = []
    print("Her bir metin parçası için özetler oluşturuluyor...")
    for i, chunk in enumerate(chunks):
        print(f"Parça {i + 1}/{len(chunks)} işleniyor...")
        prompt_chunk_summary = f"""
        Aşağıdaki mülakat metni parçasını oku. Bu parçadan, adayın aşağıda listelenen yetkinlikleri ile ilgili
        tüm önemli bilgileri, kilit ifadeleri ve somut örnekleri Türkçe olarak maddeler halinde özetle.
        Sadece bu metin parçasında geçen bilgileri kullan.

        Yetkinlikler:
        - İletişim Becerisi
        - Motivasyon ve Tutku
        - Kültürel Uyum
        - Analitik/Düşünsel Beceriler
        - Profesyonel Tutum
        - Geçmiş Deneyim Uyumu
        - Liderlik ve Girişimcilik
        - Zayıflıklarla Başa Çıkma Yetisi
        - Uzun Vadeli Potansiyel
        - Genel Etki / İzlenim

        --- MÜLAKAT METNİ PARÇASI ---
        {chunk}
        --- ÖZETİNİ BURAYA BAŞLAT ---
        """
        summary = get_llm_analysis(prompt_chunk_summary, MODEL_NAME)
        if summary:
            chunk_summaries.append(summary)
        else:
            print(f"Parça {i + 1} için özet alınamadı. İşlem durduruluyor.")
            return None, None

    combined_summary = "\n\n---\n\n".join(chunk_summaries)
    print("\nTüm parçaların özetleri başarıyla birleştirildi. Nihai analiz başlıyor.")

    print("Nihai 'Aday Değerlendirme Puanlama Tablosu' oluşturuluyor...")
    prompt_scoring = PROMPT_SCORING_DETAILS.format(candidate_name=candidate_name)
    prompt_scoring = f"""
    Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece "1. Aday Değerlendirme Puanlama Tablosu" başlığı altında yapılandır ve detaylandır:
    {prompt_scoring}

    --- MÜLAKAT ÖZETİ METNİ ---
    {combined_summary}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_scoring = get_llm_analysis(prompt_scoring, MODEL_NAME)
    if not analysis_scoring:
        print("Puanlama tablosu analizi oluşturulamadı. İşlem durduruluyor.")
        return None, None

    print("Nihai 'Recruiter Notu' oluşturuluyor...")
    prompt_recruiter = PROMPT_RECRUITER_DETAILS.format(candidate_name=candidate_name)
    prompt_recruiter = f"""
    Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece "2. Recruiter Notu" başlığı altında yapılandır ve detaylandır:
    {prompt_recruiter}

    --- MÜLAKAT ÖZETİ METNİ ---
    {combined_summary}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_recruiter = get_llm_analysis(prompt_recruiter, MODEL_NAME)
    if not analysis_recruiter:
        print("Recruiter notu analizi oluşturulamadı. İşlem durduruluyor.")
        return None, None

    return analysis_scoring, analysis_recruiter

# ==============================================================================
# --- ANA AKIŞ FONKSİYONLARI ---
# ==============================================================================

def adim_1_videodan_metne_cevir():
    """
    ADIM 1: Videodan sesi çıkarır, ElevenLabs ile deşifre eder, duygu analizi yapar ve sonucu DOCX'e kaydeder.
    """
    print("--- ADIM 1: Mülakat Videosu Metne Çevriliyor ve Duygu Analizi Yapılıyor ---")

    # 1.1. Video dosyasının varlığını kontrol et
    if not os.path.exists(INPUT_VIDEO_FILE):
        print(f"HATA: '{INPUT_VIDEO_FILE}' video dosyası bulunamadı. Lütfen kontrol edin.")
        return None, None, None

    # 1.2. Videodan sesi çıkar ve süreyi hesapla
    try:
        print(f"'{INPUT_VIDEO_FILE}' videosundan ses çıkarılıyor...")
        with VideoFileClip(INPUT_VIDEO_FILE) as video:
            video_duration_minutes = video.duration / 60
            video.audio.write_audiofile(TEMP_AUDIO_FILE)
        print(f"Ses başarıyla '{TEMP_AUDIO_FILE}' olarak kaydedildi.")
        print(f"Video süresi: {video_duration_minutes:.2f} dakika")
    except Exception as e:
        print(f"HATA: Video işlenirken bir hata oluştu: {e}")
        return None, None, None

    # 1.3. ElevenLabs ile sesi yazıya çevir
    words_data = []
    try:
        print(f"'{TEMP_AUDIO_FILE}' dosyası ElevenLabs API'sine gönderiliyor (diarize=True)...")
        elevenlabs_client = ElevenLabs(api_key=ELEVENLABS_API_KEY)
        with open(TEMP_AUDIO_FILE, "rb") as audio_file:
            response = elevenlabs_client.speech_to_text.convert(
                file=audio_file,
                model_id="scribe_v1",
                diarize=True,
            )
            words_data = response.words
        print("Ses, ElevenLabs tarafından başarıyla deşifre edildi.")
    except Exception as e:
        print(f"HATA: ElevenLabs STT işlemi sırasında hata: {e}")
        return None, None, None
    finally:
        if os.path.exists(TEMP_AUDIO_FILE):
            os.remove(TEMP_AUDIO_FILE)
            print(f"Geçici ses dosyası '{TEMP_AUDIO_FILE}' silindi.")

    # 1.4. Videodan duygu analizi yap
    duygu_cizelgesi = videodaki_duygulari_analiz_et(INPUT_VIDEO_FILE, SANIYEDE_ANALIZ_SAYISI)
    if not duygu_cizelgesi:
        print("Uyarı: Duygu analizi yapılamadı. Transkript duygu bilgisi olmadan devam edecek.")

    # 1.5. Metni konuşmacılara göre ayır ve duygu bilgisi ekle
    if not words_data:
        print("HATA: ElevenLabs'tan kelime bazlı veri alınamadı. İşlem durduruluyor.")
        return None, None, None

    diarized_output = konusmacilari_ayir_ve_cumlele(words_data, duygu_cizelgesi)
    if not diarized_output:
        print("HATA: Konuşmacı ayrımı yapılamadı veya boş sonuç döndü.")
        return None, None, None

    # 1.6. Sonucu Word dosyasına yaz
    try:
        doc = Document()
        doc.add_heading('Mülakat Transkripti', level=1)
        for entry in diarized_output:
            start_min = int(entry['baslangic'] // 60)
            start_sec = int(entry['baslangic'] % 60)
            line = f"[{entry['konusmaci']}][{start_min}:{start_sec:02d}][{entry['duygu']}]: {entry['diyalog']}"
            doc.add_paragraph(line)
        doc.save(TRANSCRIPT_DOCX_FILE)
        print(f"Transkript başarıyla '{TRANSCRIPT_DOCX_FILE}' dosyasına kaydedildi.")
        print("--- ADIM 1 TAMAMLANDI ---\n")
        return TRANSCRIPT_DOCX_FILE, video_duration_minutes, duygu_cizelgesi
    except Exception as e:
        print(f"HATA: Transkript Word dosyasına yazılırken hata oluştu: {e}")
        return None, None, None

# ==============================================================================
# --- BETİĞİ ÇALIŞTIR ---
# ==============================================================================

if __name__ == "__main__":
    print("===== MÜLAKAT VİDEOSU ANALİZ SÜRECİ BAŞLATILDI =====")

    # Adım 1: Videodan transkript ve duygu analizi
    transcript_file, video_duration, duygu_cizelgesi = adim_1_videodan_metne_cevir()

    if transcript_file and video_duration is not None:
        # Adım 2: Transkriptten adayın ismini çıkar
        candidate_name = extract_candidate_name_from_text(transcript_file)
        if not candidate_name:
            candidate_name = "Aday"
            print("Uyarı: Metinden isim tespit edilemedi, varsayılan isim 'Aday' kullanılacak.")

        # Adım 3: Videodan isimle eşleşen yüzü kırp
        final_name, face_image_path = extract_frame_and_name(INPUT_VIDEO_FILE, candidate_name, max_duration=60.0)

        # Adım 4: Görüntü tabanlı karakter analizi
        image_analysis = None
        if face_image_path and final_name == candidate_name:
            image_analysis = analyze_character_from_image(face_image_path)
            print("Adayın görüntü tabanlı karakter analizi tamamlandı.")
        else:
            print(f"Uyarı: {candidate_name} için uygun yüz görüntüsü veya eşleşen isim bulunamadı, görüntü analizi atlanıyor.")

        # Adım 5: Transkript analizi
        if video_duration > DURATION_THRESHOLD:
            print(f"Video süresi {video_duration:.2f} dakika, chunk tabanlı analiz kullanılıyor.")
            analysis_scoring, analysis_recruiter = adim_2_metin_analizi_chunk(transcript_file, candidate_name)
        else:
            print(f"Video süresi {video_duration:.2f} dakika, standart analiz kullanılıyor.")
            analysis_scoring, analysis_recruiter = adim_2_metin_analizi_yap(transcript_file, candidate_name)

        # Adım 6: Analiz sonuçlarını birleştir ve kaydet
        if analysis_scoring and analysis_recruiter:
            write_analysis_to_docx(FINAL_ANALYSIS_DOCX_FILE, analysis_scoring, analysis_recruiter, image_analysis, candidate_name)
    else:
        print("\nSüreç, Adım 1'deki bir hata nedeniyle durduruldu. Analiz yapılamadı.")

    # Geçici dosyaları temizle
    for temp_file in [TEMP_FRAME_FILE, TEMP_FACE_FILE]:
        if os.path.exists(temp_file):
            os.remove(temp_file)
            print(f"Geçici dosya '{temp_file}' silindi.")

    print("\n===== SÜREÇ TAMAMLANDI =====")