# -*- coding: utf-8 -*-
import requests
from docx import Document
import os
import json

# --- KONFİGÜRASYON ---
# Yerel LM Studio sunucunuzun çalıştığından emin olun.
LM_STUDIO_API_URL = "http://localhost:1234/v1/chat/completions"
# İndirdiğiniz GGUF model dosyasının adını doğrulayın.
MODEL_NAME = "mistral-7b-instruct-v0.3.Q4_0"
# Mülakat transkriptini içeren girdi DOCX dosyası.
INPUT_DOCX_FILE = "duygu_analizli_transkript.docx"
# Analizin kaydedileceği çıktı DOCX dosyası.
OUTPUT_DOCX_FILE = "analiz_sonucu1_mistral.docx"


def read_text_from_docx(file_path):
    """
    Bir .docx dosyasındaki tüm metni okur.
    Her paragraf okunur ve bir yeni satır karakteriyle birleştirilir.
    """
    try:
        doc = Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        return '\n'.join(full_text)
    except Exception as e:
        # Dosya bulunamadı veya bozuk docx gibi hataları yakalar.
        print(f"HATA: '{file_path}' dosyası okunurken bir hata oluştu: {e}")
        return None


def get_llm_analysis(prompt, model_name):
    """
    LM Studio API aracılığıyla yerel LLM'e bir prompt gönderir ve analizi alır.
    """
    payload = {
        "model": model_name,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7,
        "max_tokens": 1500
    }
    try:
        response = requests.post(LM_STUDIO_API_URL, json=payload)
        response.raise_for_status()  # Kötü durum kodları (4xx veya 5xx) için bir istisna oluşturur
        response_data = response.json()

        # Hata ayıklama: Tam API yanıtını yazdır
        print("API Yanıtı:", json.dumps(response_data, indent=2, ensure_ascii=False))

        if "choices" in response_data and response_data["choices"]:
            return response_data["choices"][0]["message"]["content"]
        else:
            print(f"HATA: API yanıtında 'choices' anahtarı bulunamadı veya boş. Yanıt: {response_data}")
            return None
    except requests.exceptions.ConnectionError:
        print(
            "HATA: LM Studio API sunucusuna bağlanılamadı. Lütfen LM Studio'nun çalıştığından ve sunucuyu başlattığınızdan emin olun.")
        return None
    except requests.exceptions.HTTPError as http_err:
        print(f"HATA: API isteği sırasında HTTP hatası: {http_err}")
        print(f"Yanıt İçeriği: {response.text}")
        return None
    except Exception as e:
        print(f"HATA: Analiz sırasında beklenmedik bir hata oluştu: {e}")
        return None


def analyze_interview():
    """
    Mülakat analizi sürecini yöneten ana fonksiyon.
    Transkripti bir DOCX dosyasından okur, LLM'e iki ayrı sorgu gönderir
    ve birleştirilmiş sonuçları tek bir DOCX dosyasına kaydeder.
    """
    print(f"'{INPUT_DOCX_FILE}' dosyasındaki mülakat metni okunuyor...")

    # 1. Mülakat transkriptini DOCX dosyasından oku
    interview_text = read_text_from_docx(INPUT_DOCX_FILE)
    if not interview_text or not interview_text.strip():
        print(f"HATA: '{INPUT_DOCX_FILE}' dosyası bulunamadı, bozuk veya içeriği boş. Lütfen kontrol edin.")
        return

    print("Mülakat metni başarıyla okundu.")

    # 2. Puanlama Tablosu için ilk prompt'u hazırla ve gönder
    print(f"'{MODEL_NAME}' modeline 'Aday Değerlendirme Puanlama Tablosu' için istek gönderiliyor...")
    prompt_scoring = f"""
    Aşağıdaki mülakat metnini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et. 
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:

    1. Aday Değerlendirme Puanlama Tablosu
    Adayın performansını aşağıdaki kriterlere göre 1-5 arasında puanla ve her bir puan için kısa bir gerekçe yaz. Puanlar 4/5 formatında verilsin (örneğin: 4/5), gerekçe ile birlikte ayrı bir satırda metin olarak sun (örneğin: • Kriter: 4/5 - Gerekçe). Tüm kriterlerin ortalamasını alarak Genel Ortalama Puan hesapla ve bu ortalamaya dayalı olarak İK için genel bir yorum yaz:
    • İletişim Becerisi: (1-5)
    • Motivasyon ve Tutku: (1-5)
    • Kültürel Uyum: (1-5)
    • Analitik/Düşünsel Beceriler: (1-5)
    • Profesyonel Tutum: (1-5)
    • Geçmiş Deneyim Uyumu: (1-5)
    • Liderlik ve Girişimcilik: (1-5)
    • Zayıflıklarla Başa Çıkma Yetisi: (1-5)
    • Uzun Vadeli Potansiyel: (1-5)
    • Genel Etki / İzlenim: (1-5)
    Genel Ortalama Puan: (1-5 arası hesaplanmalı)
    İK Genel Yorum: Ortalama puana dayalı genel bir değerlendirme.

    --- MÜLAKAT METNİ ---
    {interview_text}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_scoring = get_llm_analysis(prompt_scoring, MODEL_NAME)
    if not analysis_scoring:
        print("Puanlama analizi alınamadı. İşlem durduruluyor.")
        return
    print("Puanlama analizi başarıyla tamamlandı.")

    # 3. Recruiter Notu için ikinci prompt'u hazırla ve gönder
    print(f"'{MODEL_NAME}' modeline 'Recruiter Notu' için istek gönderiliyor...")
    prompt_recruiter = f"""
    Aşağıdaki mülakat metnini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:

    2. Recruiter Notu
    Adayın adını, başvurduğu pozisyonu ve genel bir yorumu içerir. Aşağıdaki alt başlıkları da ekle:
    • Aday Adı: (ad belirle)
    • Pozisyon: (pozisyon belirle)
    • Genel Yorum: Adayın genel performansı ve uygunluğu hakkında özet.
    • Dikkat Çeken Güçlü Yönler: Adayın öne çıkan becerileri veya özellikleri.
    • Geliştirme Alanları: Adayın iyileştirmesi gereken yönler.
    • Değerlendirme Önerisi: Adayın bir sonraki adımı için öneriler (ör. ikinci görüşme, teknik test).

    --- MÜLAKAT METNİ ---
    {interview_text}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_recruiter = get_llm_analysis(prompt_recruiter, MODEL_NAME)
    if not analysis_recruiter:
        print("Recruiter notu analizi alınamadı. İşlem durduruluyor.")
        return
    print("Recruiter notu analizi başarıyla tamamlandı.")

    # 4. Birleştirilmiş analiz sonuçlarını tek bir DOCX dosyasına yaz
    try:
        doc = Document()
        doc.add_heading('Mülakat Analizi Sonucu', level=0)

        # Puanlama Tablosu analizini ekle
        doc.add_heading('1. Aday Değerlendirme Puanlama Tablosu', level=1)
        # Yanıtı satırlara böl ve her satırı bir paragraf olarak ekle
        for line in analysis_scoring.split('\n'):
            if line.strip():  # Boş satır eklemekten kaçın
                doc.add_paragraph(line.strip())

        doc.add_paragraph()  # Boşluk için boş bir satır ekle

        # Recruiter Notu analizini ekle
        doc.add_heading('2. Recruiter Notu', level=1)
        for line in analysis_recruiter.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

        doc.save(OUTPUT_DOCX_FILE)
        print(f"Analiz sonuçları '{OUTPUT_DOCX_FILE}' dosyasına başarıyla kaydedildi.")
        # Kaydettikten sonra dosyayı otomatik olarak aç (isteğe bağlı)
        # os.startfile(OUTPUT_DOCX_FILE)

    except Exception as e:
        print(f"Sonuçlar dosyaya yazılırken bir hata oluştu: {e}")
        # Bir hata durumunda, ham metni hata ayıklama için kaydet
        with open("hata_ayiklama_cikti.txt", "w", encoding='utf-8') as f:
            f.write("--- PUANLAMA ANALİZİ ---\n")
            f.write(analysis_scoring)
            f.write("\n\n--- RECRUITER NOTU ---\n")
            f.write(analysis_recruiter)
        print("Ham analiz çıktıları 'hata_ayiklama_cikti.txt' dosyasına kaydedildi.")


# Ana fonksiyonu çalıştır
if __name__ == "__main__":
    analyze_interview()