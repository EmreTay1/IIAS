# Gerekli kütüphaneleri içe aktarıyoruz
import google.generativeai as genai
import os
import docx  # .docx dosyaları oluşturmak için bu kütüphaneyi ekledik

# --- KONFİGÜRASYON ---
# Google AI Studio'dan aldığınız API anahtarınızı buraya yapıştırın.
# Güvenlik için API anahtarını doğrudan koda yazmak yerine ortam değişkeni olarak ayarlamak daha iyidir.
# Ancak bu basit örnek için doğrudan kullanıyoruz.
API_KEY = "AIzaSyAyjFX1qokUL0EJ41LTjZZS7jJmq87cxjM"  # Lütfen buraya kendi API anahtarınızı girin

# Dosya isimlerini belirliyoruz
GIRIS_DOSYASI = "mulakat.txt"
# *** DEĞİŞİKLİK: Çıktı dosyası .docx olarak güncellendi. ***
CIKIS_DOSYASI = "analiz_sonucu.docx"

# --- API AYARLARI ---
# API anahtarımızı kullanarak Gemini'ı yapılandırıyoruz
try:
    genai.configure(api_key=API_KEY)
except Exception as e:
    print(f"API anahtarı yapılandırma hatası: {e}")
    print("Lütfen 'YOUR_API_KEY' yazan yere geçerli bir anahtar girdiğinizden emin olun.")
    exit()

# Kullanacağımız modeli seçiyoruz
model = genai.GenerativeModel('gemini-1.5-flash')
def mulakat_analizi_yap():
    """
    Mülakat metnini dosyadan okur, Gemini API'sine gönderir,
    gelen analizi alır ve sonuçları biçimlendirilmiş bir .docx dosyasına yazar.
    """
    print(f"'{GIRIS_DOSYASI}' dosyasındaki mülakat metni okunuyor...")

    # 1. Mülakat metnini dosyadan oku
    try:
        with open(GIRIS_DOSYASI, 'r', encoding='utf-8') as f:
            mulakat_metni = f.read()
    except FileNotFoundError:
        print(f"HATA: '{GIRIS_DOSYASI}' adında bir dosya bulunamadı.")
        print("Lütfen mülakat metnini içeren bu dosyayı kodla aynı klasöre oluşturun.")
        return
    except Exception as e:
        print(f"Dosya okunurken bir hata oluştu: {e}")
        return

    if not mulakat_metni.strip():
        print(f"HATA: '{GIRIS_DOSYASI}' dosyası boş. Lütfen mülakat metnini ekleyin.")
        return

    print("Mülakat metni başarıyla okundu.")
    print(f"'{model.model_name}' modeline gönderilerek metin analizi isteniyor...")

    # 2. Gemini API için bir prompt (istek metni) oluştur
    prompt = f"""
Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.

BÖLÜM 1: PUANLAMA TABLOSU
• Her kriteri, madde imi (•) ile başlayan ayrı bir satırda ve şu formatta yaz:
• Kriter Adı: (Puan/5) - adayin'in [puanın nedenini açıklayan kısa ve tanımlayıcı bir cümle].
• Değerlendirilecek Kriterler:
• İletişim Becerisi
• Motivasyon ve Tutku
• Kültürel Uyum
• Analitik/Düşünsel Beceriler
• Profesyonel Tutum
• Geçmiş Deneyim Uyumu
• Liderlik ve Girişimcilik
• Zayıflıklarla Başa Çıkma Yetisi
• Uzun Vadeli Potansiyel
• Genel Etki / İzlenim
• Analiz Sonu:
• Genel Ortalama Puan: Tüm puanların ortalamasını, ondalık ayraç olarak virgül kullanarak hesapla. Örnek: (3,86/5)
• İK Genel Yorum: Adayın genel potansiyelini ve ana bulguları özetleyen birkaç cümlelik bir paragraf yaz.
    --- MÜLAKAT METNİ ---
    {mulakat_metni}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """

    # 3. API'ye isteği gönder ve cevabı al
    try:
        response = model.generate_content(prompt)
        analiz_sonucu = response.text
        print("Analiz başarıyla tamamlandı.")
    except Exception as e:
        print(f"API'den cevap alınırken bir hata oluştu: {e}")
        return

    # 4. *** YENİ BÖLÜM: Gelen analizi .docx dosyasına yaz ***
    try:
        # Yeni bir Word dokümanı oluştur
        document = docx.Document()
        document.add_heading('Mülakat Analiz Raporu', level=0)

        # API'den gelen metni satır satır işle
        for line in analiz_sonucu.split('\n'):
            # Satır boş değilse devam et
            if line.strip():
                # Eğer satır bir başlık ise (ör: *1. Genel Özet:*)
                if line.strip().startswith('*') and line.strip().endswith('*'):
                    # Başlıktaki '*' karakterlerini temizle ve başlık olarak ekle (Heading 1 stili)
                    heading_text = line.replace('*', '').strip()
                    document.add_heading(heading_text, level=1)
                else:
                    # Normal bir metin satırı ise paragraf olarak ekle
                    document.add_paragraph(line.strip())

        # Hazırlanan dokümanı kaydet
        document.save(CIKIS_DOSYASI)
        print(f"Analiz sonuçları '{CIKIS_DOSYASI}' dosyasına başarıyla kaydedildi.")

    except Exception as e:
        print(f"Sonuçlar .docx dosyasına yazılırken bir hata oluştu: {e}")


# Ana fonksiyonu çalıştır
if __name__ == "__main__":
    mulakat_analizi_yap()