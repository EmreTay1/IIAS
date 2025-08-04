# -*- coding: utf-8 -*-
import requests
from docx import Document
import os
import json
from langchain.text_splitter import RecursiveCharacterTextSplitter

# --- KONFİGÜRASYON ---
LM_STUDIO_API_URL = "http://localhost:1234/v1/chat/completions"
MODEL_NAME = "mistral-7b-instruct-v0.3.Q4_0"
INPUT_DOCX_FILE = "duygu_analizli_transkript.docx"
OUTPUT_DOCX_FILE = "analiz_sonucu1_mistrail.docx"

# --- YENİ PROMPT'LAR ---
# Prompt'lar, sağlanan örnek .docx çıktısına göre tamamen yeniden düzenlendi.

PROMPT_SCORING_DETAILS = """
Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.

**BÖLÜM 1: PUANLAMA TABLOSU TALİMATLARI**

1.  **Formatlama Kuralı:** Her kriteri, madde imi (`•`) ile başlayan ayrı bir satır olarak ve aşağıdaki formatı **birebir** kullanarak yaz:
    `• Kriter Adı: (Puan/5) - Gerekçe.`

2.  **Gerekçe İçeriği:** Gerekçe bölümünde, puanın nedenini açıklayan, adayın adı ("Can Bey" gibi) ile başlayan kısa ve tanımlayıcı bir cümle kur. Örneğin: `Can Bey'in sorulara verdiği cevapların netliği.`

3.  **Değerlendirilecek Kriterler:**
    * İletişim Becerisi
    * Motivasyon ve Tutku
    * Kültürel Uyum
    * Analitik/Düşünsel Beceriler
    * Profesyonel Tutum
    * Geçmiş Deneyim Uyumu
    * Liderlik ve Girişimcilik
    * Zayıflıklarla Başa Çıkma Yetisi
    * Uzun Vadeli Potansiyel
    * Genel Etki / İzlenim

4.  **Analiz Sonu:** Puanlama listesinin ardından, aşağıdaki iki başlığı ekle:
    * `Genel Ortalama Puan:` Tüm puanların ortalamasını ondalık ayraç olarak virgül kullanarak hesapla. Örnek: `(3,86/5)`
    * `İK Genel Yorum:` Birkaç cümlelik, adayın genel potansiyelini ve ana bulguları özetleyen bir paragraf yaz.
"""

PROMPT_RECRUITER_DETAILS = """
Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.

**BÖLÜM 2: RECRUITER NOTU TALİMATLARI**

1.  -Yapısal Kural: Analizini, aşağıdaki altı başlığın tamamını madde imi (`•`) ile başlayan ayrı satırlar olarak yapılandır.

2.  **İçerik Derinliği:** Her başlığın altına, mülakat özetinden çıkardığın somut bilgilere (projeler, deneyimler, yetenekler) dayanarak detaylı ve profesyonel bir metin yaz. Adayın adını metin içinde uygun yerlerde kullan.

3.  **Başlıklar ve İçerikleri:**
    * `• Aday Adı:` Mülakat metninden adayın adını çıkar.
    * `• Pozisyon:` Mülakat metninden adayın başvurduğu pozisyonu belirle.
    * `• Genel Yorum:` Adayın geçmişi, deneyim süresi ve genel performansı hakkında özet bir paragraf yaz.
    * `• Dikkat Çeken Güçlü Yönler:` Adayın öne çıkan teknik veya sosyal yeteneklerini, projelerden örnekler vererek anlat.
    * `• Geliştirme Alanları:` Adayın hangi konularda kendini geliştirebileceğini ve potansiyel gelişim alanlarını belirt.
    * `• Değerlendirme Önerisi:` Aday için bir sonraki adımları (ikinci görüşme, teknik test vb.) ve gelişimini destekleyecek önerileri (eğitim, kurs vb.) içeren bir paragraf yaz.
"""


def read_text_from_docx(file_path):
    """
    Bir .docx dosyasındaki tüm metni okur.
    Her paragraf okunur ve bir yeni satır karakteriyle birleştirilir.
    """
    try:
        doc = Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        return '\n'.join(full_text)
    except Exception as e:
        print(f"HATA: '{file_path}' dosyası okunurken bir hata oluştu: {e}")
        return None


def write_analysis_to_docx(file_path, analysis_scoring, analysis_recruiter):
    """
    Analiz sonuçlarını yeni bir .docx dosyasına yazar.
    """
    doc = Document()
    doc.add_heading('Mülakat Analizi Sonucu', level=1)

    # Birinci kısım: Puanlama Tablosu
    doc.add_heading('1. Aday Değerlendirme Puanlama Tablosu', level=2)
    doc.add_paragraph(analysis_scoring)

    # İkinci kısım: Recruiter Notu
    doc.add_heading('2. Recruiter Notu', level=2)
    doc.add_paragraph(analysis_recruiter)

    try:
        doc.save(file_path)
        print(f"Analiz başarıyla '{file_path}' dosyasına kaydedildi.")
    except Exception as e:
        print(f"HATA: Dosyaya yazma işlemi sırasında bir hata oluştu: {e}")


def get_llm_analysis(prompt, model_name):
    """
    LM Studio API aracılığıyla yerel LLM'e bir prompt gönderir ve analizi alır.
    """
    payload = {
        "model": model_name,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7,
        "max_tokens": 2048
    }
    try:
        response = requests.post(LM_STUDIO_API_URL, json=payload)
        response.raise_for_status()
        response_data = response.json()
        if "choices" in response_data and response_data["choices"]:
            return response_data["choices"][0]["message"]["content"]
        else:
            print(f"HATA: API yanıtında 'choices' anahtarı bulunamadı veya boş. Yanıt: {response_data}")
            return None
    except requests.exceptions.ConnectionError:
        print("HATA: LM Studio API sunucusuna bağlanılamadı. Lütfen sunucunun çalıştığından emin olun.")
        return None
    except requests.exceptions.HTTPError as http_err:
        print(f"HATA: API isteği sırasında HTTP hatası: {http_err}")
        print(f"Yanıt İçeriği: {response.text}")
        return None
    except Exception as e:
        print(f"HATA: Analiz sırasında beklenmedik bir hata oluştu: {e}")
        return None


def analyze_interview_chunked():
    """
    Mülakat analizi sürecini "Böl ve Yönet" yöntemiyle yönetir.
    """
    print(f"'{INPUT_DOCX_FILE}' dosyasındaki mülakat metni okunuyor...")
    interview_text = read_text_from_docx(INPUT_DOCX_FILE)
    if not interview_text or not interview_text.strip():
        print(f"HATA: '{INPUT_DOCX_FILE}' dosyası boş veya okunamadı.")
        return

    print("Mülakat metni başarıyla okundu. Metin parçalara ayrılıyor...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=4000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(interview_text)
    print(f"Metin {len(chunks)} parçaya bölündü.")

    chunk_summaries = []
    print("Her bir metin parçası için özetler oluşturuluyor...")
    for i, chunk in enumerate(chunks):
        print(f"Parça {i + 1}/{len(chunks)} işleniyor...")
        prompt_chunk_summary = f"""
        Aşağıdaki mülakat metni parçasını oku. Bu parçadan, adayın aşağıda listelenen yetkinlikleri ile ilgili
        tüm önemli bilgileri, kilit ifadeleri ve somut örnekleri Türkçe olarak maddeler halinde özetle.
        Sadece bu metin parçasında geçen bilgileri kullan.

        Yetkinlikler:
        - İletişim Becerisi
        - Motivasyon ve Tutku
        - Kültürel Uyum
        - Analitik/Düşünsel Beceriler
        - Profesyonel Tutum
        - Geçmiş Deneyim Uyumu
        - Liderlik ve Girişimcilik
        - Zayıflıklarla Başa Çıkma Yetisi
        - Uzun Vadeli Potansiyel
        - Genel Etki / İzlenim

        --- MÜLAKAT METNİ PARÇASI ---
        {chunk}
        --- ÖZETİNİ BURAYA BAŞLAT ---
        """
        summary = get_llm_analysis(prompt_chunk_summary, MODEL_NAME)
        if summary:
            chunk_summaries.append(summary)
        else:
            print(f"Parça {i + 1} için özet alınamadı. İşlem durduruluyor.")
            return

    combined_summary = "\n\n---\n\n".join(chunk_summaries)
    print("\nTüm parçaların özetleri başarıyla birleştirildi. Nihai analiz başlıyor.")

    print("Nihai 'Aday Değerlendirme Puanlama Tablosu' oluşturuluyor...")
    prompt_scoring = f"""
    Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece "1. Aday Değerlendirme Puanlama Tablosu" başlığı altında yapılandır ve detaylandır:
    {PROMPT_SCORING_DETAILS}

    --- MÜLAKAT ÖZETİ METNİ ---
    {combined_summary}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_scoring = get_llm_analysis(prompt_scoring, MODEL_NAME)

    if not analysis_scoring:
        print("Puanlama tablosu analizi oluşturulamadı. İşlem durduruluyor.")
        return

    print("Nihai 'Recruiter Notu' oluşturuluyor...")
    prompt_recruiter = f"""
    Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece "2. Recruiter Notu" başlığı altında yapılandır ve detaylandır:
    {PROMPT_RECRUITER_DETAILS}

    --- MÜLAKAT ÖZETİ METNİ ---
    {combined_summary}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_recruiter = get_llm_analysis(prompt_recruiter, MODEL_NAME)

    if not analysis_recruiter:
        print("Recruiter notu analizi oluşturulamadı. İşlem durduruluyor.")
        return

    print("Analiz sonuçları birleştirilerek DOCX dosyasına yazılıyor...")
    write_analysis_to_docx(OUTPUT_DOCX_FILE, analysis_scoring, analysis_recruiter)


# Ana fonksiyonu çalıştır
if __name__ == "__main__":
    analyze_interview_chunked()