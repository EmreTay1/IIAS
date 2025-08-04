# Gerekli kütüphaneleri içe aktarıyoruz
from mistralai import Mistral
from docx import Document
import os

# --- KONFİGÜRASYON ---
API_KEY = "fuayrW8itfw63EIEj5I6lqb0n4MWz5JK"  # Mistral AI'dan aldığınız API anahtarını buraya girin

# Dosya isimlerini belirliyoruz
GIRIS_DOSYASI = "mulakat.txt"
CIKIS_DOSYASI = "analiz_sonucu_mistral.docx"

# --- API AYARLARI ---
try:
    # Kullanacağımız modeli seçiyoruz
    MODEL_ADI = "mistral-small-latest"
    client = Mistral(api_key=API_KEY)
except Exception as e:
    print(f"API anahtarı veya istemci yapılandırma hatası: {e}")
    print("Lütfen geçerli bir API anahtarı girdiğinizden emin olun.")
    exit()

def mulakat_analizi_yap():
    """
    Mülakat metnini dosyadan okur, Mistral AI API'sine gönderir,
    gelen analizi alır ve sonuçları bir .docx dosyasına yazar.
    """
    print(f"'{GIRIS_DOSYASI}' dosyasındaki mülakat metni okunuyor...")

    # 1. Mülakat metnini dosyadan oku
    try:
        with open(GIRIS_DOSYASI, 'r', encoding='utf-8') as f:
            mulakat_metni = f.read()
    except FileNotFoundError:
        print(f"HATA: '{GIRIS_DOSYASI}' adında bir dosya bulunamadı.")
        print("Lütfen mülakat metnini içeren bu dosyayı kodla aynı klasöre oluşturun.")
        return
    except Exception as e:
        print(f"Dosya okunurken bir hata oluştu: {e}")
        return

    if not mulakat_metni.strip():
        print(f"HATA: '{GIRIS_DOSYASI}' dosyası boş. Lütfen mülakat metnini ekleyin.")
        return

    print("Mülakat metni başarıyla okundu.")
    print(f"'{MODEL_ADI}' modeline gönderilerek metin analizi isteniyor...")

    # 2. Prompt oluştur
    prompt = f"""
     Aşağıdaki mülakat metnini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et. 
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:

    1. Aday Değerlendirme Puanlama Tablosu
    Adayın performansını aşağıdaki kriterlere göre 1-5 arasında puanla ve her bir puan için kısa bir gerekçe yaz. Puanlar 4/5 formatında verilsin (örneğin: 4/5), gerekçe ile birlikte ayrı bir satırda metin olarak sun (örneğin: • Kriter: 4/5 - Gerekçe). Tüm kriterlerin ortalamasını alarak Genel Ortalama Puan hesapla ve bu ortalamaya dayalı olarak İK için genel bir yorum yaz:
    • İletişim Becerisi: (1-5)
    • Motivasyon ve Tutku: (1-5)
    • Kültürel Uyum: (1-5)
    • Analitik/Düşünsel Beceriler: (1-5)
    • Profesyonel Tutum: (1-5)
    • Geçmiş Deneyim Uyumu: (1-5)
    • Liderlik ve Girişimcilik: (1-5)
    • Zayıflıklarla Başa Çıkma Yetisi: (1-5)
    • Uzun Vadeli Potansiyel: (1-5)
    • Genel Etki / İzlenim: (1-5)
    Genel Ortalama Puan: (1-5 arası hesaplanmalı)
    İK Genel Yorum: Ortalama puana dayalı genel bir değerlendirme.
    
    
    2. Recruiter Notu
    Adayın adını, başvurduğu pozisyonu ve genel bir yorumu içerir. Aşağıdaki alt başlıkları da ekle:
    • Aday Adı: (ad belirle)
    • Pozisyon: (pozisyon belirle)
    • Genel Yorum: Adayın genel performansı ve uygunluğu hakkında özet.
    • Dikkat Çeken Güçlü Yönler: Adayın öne çıkan becerileri veya özellikleri.
    • Geliştirme Alanları: Adayın iyileştirmesi gereken yönler.
    • Değerlendirme Önerisi: Adayın bir sonraki adımı için öneriler (ör. ikinci görüşme, teknik test).

    --- MÜLAKAT METNİ ---
    {mulakat_metni}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """

    # 3. API'ye isteği Mistral formatında gönder ve cevabı al
    try:
        messages = [
            {"role": "user", "content": prompt}
        ]

        chat_response = client.chat.complete(
            model=MODEL_ADI,
            messages=messages
        )

        analiz_sonucu = chat_response.choices[0].message.content
        print("Analiz başarıyla tamamlandı.")
    except Exception as e:
        print(f"API'den cevap alınırken bir hata oluştu: {e}")
        return

    # 4. Gelen analizi .docx dosyasına yaz
    try:
        doc = Document()
        doc.add_heading('Mülakat Analizi Sonucu', 0)
        doc.add_paragraph(analiz_sonucu)
        doc.save(CIKIS_DOSYASI)
        print(f"Analiz sonuçları '{CIKIS_DOSYASI}' dosyasına başarıyla kaydedildi.")
    except Exception as e:
        print(f"Sonuçlar dosyaya yazılırken bir hata oluştu: {e}")

# Ana fonksiyonu çalıştır
if __name__ == "__main__":
    mulakat_analizi_yap()