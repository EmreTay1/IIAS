# Gerekli kütüphaneleri içe aktar
from moviepy import VideoFileClip
import os
import spacy
from elevenlabs import ElevenLabs
from docx import Document  # <-- Word dosyası oluşturmak için eklendi

# --- 1. ElevenLabs API istemcisini başlat ---
# ÖNEMLİ: Gerçek API anahtarınızı tırnak işaretleri arasına yapıştırın.
# api_key'inizi korumak için bu kodu başkalarıyla paylaşırken dikkatli olun.
elevenlabs_client = ElevenLabs(
    api_key="sk_062d16b49c532f296309a61825e5c24bb4915f8586b1cdcc"  # ← BURAYA kendi API anahtarını yaz
)

# --- 2. Videodan sesi çıkar ---
video_dosyasi = "mulakat.mp4"
audio_path = "gecici_ses.wav"

# Video dosyasının var olup olmadığını kontrol et
if not os.path.exists(video_dosyasi):
    print(f"Hata: '{video_dosyasi}' adında bir video dosyası bulunamadı.")
    print("Lütfen doğru dosya adını girdiğinizden ve dosyanın kodla aynı dizinde olduğundan emin olun.")
    exit()

try:
    print(f"'{video_dosyasi}' videosundan ses çıkarılıyor...")
    video = VideoFileClip(video_dosyasi)
    video.audio.write_audiofile(audio_path)
    print(f"Ses başarıyla '{audio_path}' olarak kaydedildi.")
except Exception as e:
    print(f"Video işlenirken bir hata oluştu: {e}")
    exit()

# --- 3. ElevenLabs ile sesi yazıya çevir ---
metin = ""
words_data = []  # ElevenLabs'tan gelen kelime bazlı veriyi tutmak için
try:
    print(f"'{audio_path}' dosyası ElevenLabs API'sine gönderiliyor... (diarize=True ile)")
    with open(audio_path, "rb") as audio_file:
        result = elevenlabs_client.speech_to_text.convert(
            file=audio_file,
            model_id="scribe_v1",      # Speech to text için kullanılan model
            diarize=True,             # Konuşmacı ayırmayı etkinleştir
        )
        metin = result.text          # Genel transkripsiyonu sakla
        words_data = result.words    # Kelime bazlı diarizasyon bilgisini al

        print("\nYazıya Çevrilen Ham Metin:")
        print(metin)

except Exception as e:
    print(f"ElevenLabs STT hatası: {e}")
    metin = ""

# --- 4. Geçici ses dosyasını sil ---
if os.path.exists(audio_path):
    os.remove(audio_path)
    print(f"\nGeçici dosya '{audio_path}' silindi.")


# --- 5. Metni konuşmacılara göre ayır (ElevenLabs kelime verisi kullanarak) ---
def konusmacilari_ayir_gelismis(words_data):
    if not words_data:
        print("\nKonuşmacı ayırmak için kelime verisi bulunamadı.")
        return []

    speaker_map = {}
    next_speaker_num = 1
    current_speaker_id_raw = None
    current_speaker_mapped_label = None
    current_speaker_words_buffer = []
    all_speaker_utterances_combined = []

    for word_obj in words_data:
        word_text = word_obj.text
        raw_speaker_id = word_obj.speaker_id

        if raw_speaker_id not in speaker_map:
            speaker_map[raw_speaker_id] = f"Konuşmacı {next_speaker_num}"
            next_speaker_num += 1

        mapped_speaker_label = speaker_map[raw_speaker_id]

        if current_speaker_id_raw is None:
            current_speaker_id_raw = raw_speaker_id
            current_speaker_mapped_label = mapped_speaker_label
            current_speaker_words_buffer.append(word_text)
        elif raw_speaker_id != current_speaker_id_raw:
            if current_speaker_words_buffer:
                all_speaker_utterances_combined.append({
                    'speaker': current_speaker_mapped_label,
                    'text_combined': "".join(current_speaker_words_buffer).strip()
                })
            current_speaker_id_raw = raw_speaker_id
            current_speaker_mapped_label = mapped_speaker_label
            current_speaker_words_buffer = [word_text]
        else:
            current_speaker_words_buffer.append(word_text)

    if current_speaker_words_buffer:
        all_speaker_utterances_combined.append({
            'speaker': current_speaker_mapped_label,
            'text_combined': "".join(current_speaker_words_buffer).strip()
        })

    final_diarized_sentences = []
    try:
        nlp = spacy.load("xx_ent_wiki_sm")
        if "sentencizer" not in nlp.pipe_names:
            nlp.add_pipe("sentencizer")
    except OSError:
        print("\nspaCy modeli 'xx_ent_wiki_sm' bulunamadı.")
        print("Lütfen 'python -m spacy download xx_ent_wiki_sm' komutunu çalıştırın.")
        return []
    except Exception as e:
        print(f"spaCy modeli yüklenirken bir hata oluştu: {e}")
        return []

    for entry in all_speaker_utterances_combined:
        doc = nlp(entry['text_combined'])
        sentences_for_speaker = [sent.text.strip() for sent in doc.sents if sent.text.strip()]
        for sent in sentences_for_speaker:
            final_diarized_sentences.append({
                "konusmaci": entry['speaker'],
                "diyalog": sent
            })

    return final_diarized_sentences


# --- 6. Konuşmacılara ayırma işlemini çalıştır ---
diarized_output = []
if words_data:
    diarized_output = konusmacilari_ayir_gelismis(words_data)
else:
    print("\nElevenLabs'tan kelime bazlı diarizasyon verisi alınamadığı için ayırma işlemi yapılamadı.")


# --- 7. Sonucu Word dosyasına yaz ---
if diarized_output:
    output_filename = "mulakat.docx"
    try:
        # Yeni bir Word belgesi oluştur
        doc = Document()
        doc.add_heading('Mülakat Transkripti', level=1)
        print(f"\n--- Çıktı '{output_filename}' Dosyasına Yazılıyor ---")

        # Ayrılmış diyalogları döngüye al
        for entry in diarized_output:
            # Hem konsola yazdır (isteğe bağlı) hem de dosyaya ekle
            line = f"[{entry['konusmaci']}]: {entry['diyalog']}"
            print(line)
            doc.add_paragraph(line)

        # Belgeyi kaydet
        doc.save(output_filename)
        print(f"\nTranskript başarıyla '{output_filename}' dosyasına kaydedildi.")

    except Exception as e:
        print(f"\nWord dosyası oluşturulurken bir hata oluştu: {e}")
        print("Lütfen 'python-docx' kütüphanesinin kurulu olduğundan emin olun: pip install python-docx")

else:
    print("\nKonuşmacı ayrımı yapılamadı veya boş sonuç döndürüldü. Word dosyası oluşturulmadı.")