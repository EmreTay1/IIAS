# -*- coding: utf-8 -*-

# --- GEREKLİ KÜTÜPHANELER ---
import os
import json
import requests
import spacy
from moviepy import VideoFileClip
from docx import Document
from elevenlabs import ElevenLabs

# --- KONFİGÜRASYON ---
# Bu bölümdeki değişkenleri kendi ayarlarınıza göre güncelleyin.

# --- 1. Adım: Ses Deşifre Ayarları ---
# ElevenLabs API Anahtarınız
# ÖNEMLİ: Gerçek API anahtarınızı tırnak işaretleri arasına yapıştırın.
ELEVENLABS_API_KEY = "sk_062d16b49c532f296309a61825e5c24bb4915f8586b1cdcc"
# Girdi olarak kullanılacak video dosyası
INPUT_VIDEO_FILE = "mulakat.mp4"
# Videodan çıkarılacak geçici ses dosyasının adı
TEMP_AUDIO_FILE = "gecici_ses.wav"
# Deşifre edilen metnin kaydedileceği ara Word dosyası
TRANSCRIPT_DOCX_FILE = "mulakat_transkripti.docx"

# --- 2. Adım: Metin Analizi Ayarları ---
# Yerel LM Studio sunucunuzun API adresi.
LM_STUDIO_API_URL = "http://localhost:1234/v1/chat/completions"
# LM Studio'da kullandığınız modelin tam adı.
MODEL_NAME = "mistral-7b-instruct-v0.3.Q4_0"
# Analiz sonuçlarının kaydedileceği nihai Word dosyası.
FINAL_ANALYSIS_DOCX_FILE = "analiz_sonucu.docx"


# ==============================================================================
# --- YARDIMCI FONKSİYONLAR ---
# ==============================================================================

def konusmacilari_ayir_ve_cumlele(words_data):
    """
    ElevenLabs'ten gelen kelime bazlı diarizasyon verisini işler.
    Konuşmacı etiketlerine göre metni gruplar ve spaCy ile cümlelere ayırır.
    """
    if not words_data:
        print("\nKonuşmacı ayırmak için kelime verisi bulunamadı.")
        return []

    speaker_map = {}
    next_speaker_num = 1
    current_speaker_id_raw = None
    current_speaker_mapped_label = None
    current_speaker_words_buffer = []
    all_speaker_utterances_combined = []

    for word_obj in words_data:
        word_text = word_obj.text
        raw_speaker_id = word_obj.speaker_id

        if raw_speaker_id not in speaker_map:
            speaker_map[raw_speaker_id] = f"Konuşmacı {next_speaker_num}"
            next_speaker_num += 1

        mapped_speaker_label = speaker_map[raw_speaker_id]

        if current_speaker_id_raw is None:
            current_speaker_id_raw = raw_speaker_id
            current_speaker_mapped_label = mapped_speaker_label
        elif raw_speaker_id != current_speaker_id_raw:
            if current_speaker_words_buffer:
                all_speaker_utterances_combined.append({
                    'speaker': current_speaker_mapped_label,
                    'text_combined': "".join(current_speaker_words_buffer).strip()
                })
            current_speaker_id_raw = raw_speaker_id
            current_speaker_mapped_label = mapped_speaker_label
            current_speaker_words_buffer = []

        current_speaker_words_buffer.append(word_text)

    if current_speaker_words_buffer:
        all_speaker_utterances_combined.append({
            'speaker': current_speaker_mapped_label,
            'text_combined': "".join(current_speaker_words_buffer).strip()
        })

    final_diarized_sentences = []
    try:
        nlp = spacy.load("xx_ent_wiki_sm")
        if "sentencizer" not in nlp.pipe_names:
            nlp.add_pipe("sentencizer")
    except OSError:
        print("\nspaCy modeli 'xx_ent_wiki_sm' bulunamadı.")
        print("Lütfen 'python -m spacy download xx_ent_wiki_sm' komutunu çalıştırın.")
        return []
    except Exception as e:
        print(f"spaCy modeli yüklenirken bir hata oluştu: {e}")
        return []

    for entry in all_speaker_utterances_combined:
        doc = nlp(entry['text_combined'])
        sentences_for_speaker = [sent.text.strip() for sent in doc.sents if sent.text.strip()]
        for sent in sentences_for_speaker:
            final_diarized_sentences.append({
                "konusmaci": entry['speaker'],
                "diyalog": sent
            })

    return final_diarized_sentences


def read_text_from_docx(file_path):
    """
    Bir .docx dosyasındaki tüm metni okur ve tek bir metin bloğu olarak döndürür.
    """
    try:
        doc = Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        return '\n'.join(full_text)
    except Exception as e:
        print(f"HATA: '{file_path}' dosyası okunurken bir hata oluştu: {e}")
        return None


def get_llm_analysis(prompt, model_name):
    """
    LM Studio API aracılığıyla yerel LLM'e bir prompt gönderir ve analizi alır.
    """
    payload = {
        "model": model_name,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7,
        "max_tokens": 1500
    }
    try:
        response = requests.post(LM_STUDIO_API_URL, json=payload)
        response.raise_for_status()
        response_data = response.json()

        if "choices" in response_data and response_data["choices"]:
            return response_data["choices"][0]["message"]["content"]
        else:
            print(f"HATA: API yanıtında 'choices' anahtarı bulunamadı veya boş. Yanıt: {response_data}")
            return None
    except requests.exceptions.ConnectionError:
        print("HATA: LM Studio API sunucusuna bağlanılamadı. Lütfen LM Studio'nun çalıştığından emin olun.")
        return None
    except requests.exceptions.HTTPError as http_err:
        print(f"HATA: API isteği sırasında HTTP hatası: {http_err}\nYanıt İçeriği: {response.text}")
        return None
    except Exception as e:
        print(f"HATA: Analiz sırasında beklenmedik bir hata oluştu: {e}")
        return None


# ==============================================================================
# --- ANA AKIŞ FONKSİYONLARI ---
# ==============================================================================

def adim_1_videodan_metne_cevir():
    """
    ADIM 1: Videodan sesi çıkarır, ElevenLabs ile deşifre eder ve sonucu DOCX'e kaydeder.
    Başarılı olursa transkript dosyasının yolunu döndürür, aksi takdirde None.
    """
    print("--- ADIM 1: Mülakat Videosu Metne Çevriliyor ---")

    # 1.1. Video dosyasının varlığını kontrol et
    if not os.path.exists(INPUT_VIDEO_FILE):
        print(f"HATA: '{INPUT_VIDEO_FILE}' video dosyası bulunamadı. Lütfen kontrol edin.")
        return None

    # 1.2. Videodan sesi çıkar
    try:
        print(f"'{INPUT_VIDEO_FILE}' videosundan ses çıkarılıyor...")
        with VideoFileClip(INPUT_VIDEO_FILE) as video:
            video.audio.write_audiofile(TEMP_AUDIO_FILE)
        print(f"Ses başarıyla '{TEMP_AUDIO_FILE}' olarak kaydedildi.")
    except Exception as e:
        print(f"HATA: Video işlenirken bir hata oluştu: {e}")
        return None

    # 1.3. ElevenLabs ile sesi yazıya çevir
    words_data = []
    try:
        print(f"'{TEMP_AUDIO_FILE}' dosyası ElevenLabs API'sine gönderiliyor (diarize=True)...")
        elevenlabs_client = ElevenLabs(api_key=ELEVENLABS_API_KEY)
        with open(TEMP_AUDIO_FILE, "rb") as audio_file:
            response = elevenlabs_client.speech_to_text.convert(
                file=audio_file,
                model_id="scribe_v1",
                diarize=True,
            )
            words_data = response.words
        print("Ses, ElevenLabs tarafından başarıyla deşifre edildi.")
    except Exception as e:
        print(f"HATA: ElevenLabs STT işlemi sırasında hata: {e}")
        return None
    finally:
        # 1.4. Geçici ses dosyasını sil
        if os.path.exists(TEMP_AUDIO_FILE):
            os.remove(TEMP_AUDIO_FILE)
            print(f"Geçici ses dosyası '{TEMP_AUDIO_FILE}' silindi.")

    # 1.5. Metni konuşmacılara göre ayır
    if not words_data:
        print("HATA: ElevenLabs'tan kelime bazlı veri alınamadı. İşlem durduruluyor.")
        return None

    diarized_output = konusmacilari_ayir_ve_cumlele(words_data)
    if not diarized_output:
        print("HATA: Konuşmacı ayrımı yapılamadı veya boş sonuç döndü.")
        return None

    # 1.6. Sonucu Word dosyasına yaz
    try:
        doc = Document()
        doc.add_heading('Mülakat Transkripti', level=1)
        for entry in diarized_output:
            line = f"[{entry['konusmaci']}]: {entry['diyalog']}"
            doc.add_paragraph(line)
        doc.save(TRANSCRIPT_DOCX_FILE)
        print(f"Transkript başarıyla '{TRANSCRIPT_DOCX_FILE}' dosyasına kaydedildi.")
        print("--- ADIM 1 TAMAMLANDI ---\n")
        return TRANSCRIPT_DOCX_FILE
    except Exception as e:
        print(f"HATA: Transkript Word dosyasına yazılırken hata oluştu: {e}")
        return None


def adim_2_metin_analizi_yap(transcript_file_path):
    """
    ADIM 2: Transkripti DOCX dosyasından okur, LLM'e analiz için gönderir
    ve sonuçları nihai bir DOCX dosyasına kaydeder.
    """
    print(f"--- ADIM 2: '{transcript_file_path}' Dosyasındaki Metin Analiz Ediliyor ---")

    # 2.1. Mülakat transkriptini DOCX dosyasından oku
    interview_text = read_text_from_docx(transcript_file_path)
    if not interview_text or not interview_text.strip():
        print(f"HATA: '{transcript_file_path}' dosyası bulunamadı, bozuk veya boş. Analiz yapılamıyor.")
        return

    print("Mülakat transkripti başarıyla okundu.")

    # 2.2. Puanlama Tablosu için prompt'u hazırla ve LLM'e gönder
    print(f"'{MODEL_NAME}' modeline 'Aday Değerlendirme Puanlama Tablosu' için istek gönderiliyor...")
    prompt_scoring = f"""
   Aşağıdaki mülakat metnini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et. 
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:

    1. Aday Değerlendirme Puanlama Tablosu
    Adayın performansını aşağıdaki kriterlere göre 1-5 arasında puanla ve her bir puan için kısa bir gerekçe yaz. Puanlar 4/5 formatında verilsin (örneğin: 4/5), gerekçe ile birlikte ayrı bir satırda metin olarak sun (örneğin: • Kriter: 4/5 - Gerekçe). Tüm kriterlerin ortalamasını alarak Genel Ortalama Puan hesapla ve bu ortalamaya dayalı olarak İK için genel bir yorum yaz:
    • İletişim Becerisi: (1-5)
    • Motivasyon ve Tutku: (1-5)
    • Kültürel Uyum: (1-5)
    • Analitik/Düşünsel Beceriler: (1-5)
    • Profesyonel Tutum: (1-5)
    • Geçmiş Deneyim Uyumu: (1-5)
    • Liderlik ve Girişimcilik: (1-5)
    • Zayıflıklarla Başa Çıkma Yetisi: (1-5)
    • Uzun Vadeli Potansiyel: (1-5)
    • Genel Etki / İzlenim: (1-5)
    Genel Ortalama Puan: (1-5 arası hesaplanmalı)
    İK Genel Yorum: Ortalama puana dayalı genel bir değerlendirme.

    --- MÜLAKAT METNİ ---
    {interview_text}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_scoring = get_llm_analysis(prompt_scoring, MODEL_NAME)
    if not analysis_scoring:
        print("Puanlama analizi alınamadı. İşlem durduruluyor.")
        return
    print("Puanlama analizi başarıyla tamamlandı.")

    # 2.3. Recruiter Notu için prompt'u hazırla ve LLM'e gönder
    print(f"'{MODEL_NAME}' modeline 'Recruiter Notu' için istek gönderiliyor...")
    prompt_recruiter = f"""
    Aşağıdaki mülakat metnini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:

    2. Recruiter Notu
    Adayın adını, başvurduğu pozisyonu ve genel bir yorumu içerir. Aşağıdaki alt başlıkları da ekle:
    • Aday Adı: (ad belirle)
    • Pozisyon: (pozisyon belirle)
    • Genel Yorum: Adayın genel performansı ve uygunluğu hakkında özet.
    • Dikkat Çeken Güçlü Yönler: Adayın öne çıkan becerileri veya özellikleri.
    • Geliştirme Alanları: Adayın iyileştirmesi gereken yönler.
    • Değerlendirme Önerisi: Adayın bir sonraki adımı için öneriler (ör. ikinci görüşme, teknik test).

    --- MÜLAKAT METNİ ---
    {interview_text}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_recruiter = get_llm_analysis(prompt_recruiter, MODEL_NAME)
    if not analysis_recruiter:
        print("Recruiter notu analizi alınamadı. İşlem durduruluyor.")
        return
    print("Recruiter notu analizi başarıyla tamamlandı.")

    # 2.4. Birleştirilmiş analiz sonuçlarını tek bir DOCX dosyasına yaz
    try:
        doc = Document()
        doc.add_heading('Mülakat Analizi Sonucu', level=0)

        doc.add_heading('1. Aday Değerlendirme Puanlama Tablosu', level=1)
        for line in analysis_scoring.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

        doc.add_paragraph()

        doc.add_heading('2. Recruiter Notu', level=1)
        for line in analysis_recruiter.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

        doc.save(FINAL_ANALYSIS_DOCX_FILE)
        print(f"\nAnaliz sonuçları '{FINAL_ANALYSIS_DOCX_FILE}' dosyasına başarıyla kaydedildi.")
        print("--- ADIM 2 TAMAMLANDI ---")

    except Exception as e:
        print(f"HATA: Sonuçlar dosyaya yazılırken bir hata oluştu: {e}")


# ==============================================================================
# --- BETİĞİ ÇALIŞTIR ---
# ==============================================================================

if __name__ == "__main__":
    print("===== MÜLAKAT VİDEOSU ANALİZ SÜRECİ BAŞLATILDI =====")

    # Adım 1'i çalıştır: Videoyu metne çevir.
    transcript_file = adim_1_videodan_metne_cevir()

    # Adım 1 başarılı olduysa, Adım 2'ye geç.
    if transcript_file:
        # Adım 2'yi çalıştır: Oluşturulan metin dosyasını analiz et.
        adim_2_metin_analizi_yap(transcript_file)

        # (İsteğe bağlı) İşlem bittikten sonra ara transkript dosyasını sil
        # os.remove(transcript_file)
        # print(f"\nAra transkript dosyası '{transcript_file}' silindi.")

    else:
        print("\nSüreç, Adım 1'deki bir hata nedeniyle durduruldu. Analiz yapılamadı.")

    print("\n===== SÜREÇ TAMAMLANDI =====")