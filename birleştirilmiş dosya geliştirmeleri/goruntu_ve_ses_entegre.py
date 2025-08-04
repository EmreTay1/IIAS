# Gerekli kütüphaneleri içe aktar
import os
import cv2
import spacy
from moviepy import VideoFileClip
from elevenlabs.client import ElevenLabs
from docx import Document
from deepface import DeepFace
from collections import Counter
import warnings

# DeepFace ve TensorFlow'un başlangıçta gösterdiği teknik uyarıları gizle
warnings.filterwarnings("ignore", category=UserWarning)

# --- 1. AYARLAR VE KONFİGÜRASYON ---

# ÖNEMLİ: Gerçek API anahtarınızı tırnak işaretleri arasına yapıştırın.
# Bu anahtarı güvende tutun ve başkalarıyla paylaşmayın.
ELEVENLABS_API_KEY = "sk_791994ad759e90a722561c1f23164f6eb8ab47d84ae31d0d"

# İşlenecek video dosyasının adı
VIDEO_DOSYASI = "(İK)Mülakat - Trim.mp4"  # ← Kendi video dosyanızın adını yazın

# Çıktı olarak oluşturulacak Word dosyasının adı
CIKTI_DOSYASI = "duygu_analizli_transkript.docx"

# Analiz sırasında oluşturulacak geçici ses dosyasının yolu
GECICI_SES_DOSYASI = "temp_audio_for_analysis.wav"

# Performans ayarı: Saniyede kaç kez duygu analizi yapılacak?
# Yüksek değerler daha hassas ama yavaş, düşük değerler daha hızlı ama daha az hassas olur.
SANIYEDE_ANALIZ_SAYISI = 2

# DeepFace'in İngilizce duygu etiketlerini Türkçeye çevirmek için bir sözlük
DUYGU_SOZLUGU = {
    'angry': 'ÖFKELİ',
    'disgust': 'TİKSİNMİŞ',
    'fear': 'KORKMUŞ',
    'happy': 'MUTLU',
    'sad': 'ÜZGÜN',
    'surprise': 'ŞAŞKIN',
    'neutral': 'NÖTR'
}


# --- 2. YARDIMCI FONKSİYONLAR ---

def videodan_sesi_cikar(video_path, audio_path):
    """Videodan sesi ayıklar ve belirtilen yola kaydeder."""
    if not os.path.exists(video_path):
        print(f"Hata: '{video_path}' video dosyası bulunamadı. Lütfen dosya adını kontrol edin.")
        return False
    try:
        print(f"'{video_path}' videosundan ses çıkarılıyor...")
        video = VideoFileClip(video_path)
        video.audio.write_audiofile(audio_path, codec='pcm_s16le')
        print(f"Ses başarıyla '{audio_path}' olarak kaydedildi.")
        return True
    except Exception as e:
        print(f"Video işlenirken bir hata oluştu: {e}")
        return False


def sesi_yaziya_cevir(api_key, audio_path):
    """ElevenLabs kullanarak sesi metne çevirir ve kelime bazlı zaman damgalarını döndürür."""
    try:
        print(f"'{audio_path}' dosyası ElevenLabs API'sine gönderiliyor...")
        client = ElevenLabs(api_key=api_key)
        with open(audio_path, "rb") as audio_file:
            response = client.speech_to_text.convert(
                file=audio_file,
                model_id="scribe_v1",  # Konuşmadan metne çevirme modeli
                diarize=True  # Konuşmacı ayırmayı ve zaman damgalarını etkinleştir
            )
        print("Sesi yazıya çevirme işlemi tamamlandı.")
        if not response.words:
            print("Uyarı: ElevenLabs API'si kelime bazında zaman damgası veya konuşmacı verisi döndürmedi.")
            return None
        return response.words
    except Exception as e:
        print(f"ElevenLabs STT (Speech-to-Text) hatası: {e}")
        return None


def videodaki_duygulari_analiz_et(video_path, saniyede_kontrol=2):
    """Videoyu analiz eder ve zaman damgalarıyla birlikte dominant duyguları listeler."""
    print("\nVideo üzerinden duygu analizi başlatılıyor (bu işlem videonun uzunluğuna göre sürebilir)...")
    video = cv2.VideoCapture(video_path)
    fps = video.get(cv2.CAP_PROP_FPS)
    # Saniyede_kontrol=0 ise veya fps=0 ise hatayı önle
    frame_interval = int(fps / saniyede_kontrol) if saniyede_kontrol > 0 and fps > 0 else int(fps)

    duygu_zaman_cizelgesi = []
    frame_num = 0

    while video.isOpened():
        ret, frame = video.read()
        if not ret:
            break

        # Sadece belirli aralıklardaki kareleri analiz et
        if frame_num % frame_interval == 0:
            try:
                analysis = DeepFace.analyze(
                    frame,
                    actions=['emotion'],
                    enforce_detection=False,  # Karede yüz bulamazsa hata verme
                    silent=True  # DeepFace'in kendi ilerleme çubuğunu gizle
                )
                if isinstance(analysis, list) and len(analysis) > 0:
                    dominant_emotion_en = analysis[0]['dominant_emotion']
                    dominant_emotion_tr = DUYGU_SOZLUGU.get(dominant_emotion_en, dominant_emotion_en.upper())
                    timestamp = frame_num / fps
                    duygu_zaman_cizelgesi.append({'zaman': timestamp, 'duygu': dominant_emotion_tr})
            except Exception:
                pass  # Yüz bulunamayan veya analiz edilemeyen kareleri atla
        frame_num += 1

    video.release()
    print("Duygu analizi tamamlandı.")
    return duygu_zaman_cizelgesi


def verileri_birlestir_ve_formatla(kelime_verisi, duygu_cizelgesi, nlp_model):
    """Konuşma metnini ve duygu analizini birleştirerek son çıktıyı oluşturur."""
    if not kelime_verisi:
        return []

    # Cümleleri ve duyguları birleştirmek için ana liste
    sonuclar = []

    # Kelimeleri önce konuşmacı ve metin olarak birleştir
    mevcut_konusmaci = None
    konusma_blogu = []
    tum_konusma_bloklari = []

    for kelime in kelime_verisi:
        if mevcut_konusmaci is None:
            mevcut_konusmaci = kelime.speaker_id

        if kelime.speaker_id != mevcut_konusmaci:
            tum_konusma_bloklari.append({
                "speaker_id": mevcut_konusmaci,
                "words": konusma_blogu
            })
            mevcut_konusmaci = kelime.speaker_id
            konusma_blogu = []

        konusma_blogu.append(kelime)

    # Son bloğu da ekle
    if konusma_blogu:
        tum_konusma_bloklari.append({
            "speaker_id": mevcut_konusmaci,
            "words": konusma_blogu
        })

    # Konuşmacıları numaralandır (Konuşmacı 1, Konuşmacı 2)
    konusmaci_map = {speaker_id: f"Konuşmacı {i + 1}" for i, speaker_id in
                     enumerate(dict.fromkeys(k['speaker_id'] for k in tum_konusma_bloklari))}

    # Her konuşma bloğunu cümlelere ayır ve duygusunu bul
    for blok in tum_konusma_bloklari:
        metin = " ".join([k.text for k in blok['words']])
        doc = nlp_model(metin)

        for cumle in doc.sents:
            cumle_metni = cumle.text.strip()
            if not cumle_metni:
                continue

            # Cümledeki kelimelerin başlangıç ve bitiş zamanını bul
            cumle_kelimeleri = [k for k in blok['words'] if k.text in cumle_metni]
            if not cumle_kelimeleri:
                continue

            baslangic_zamani = min(k.start for k in cumle_kelimeleri)
            bitis_zamani = max(k.end for k in cumle_kelimeleri)

            # O zaman aralığındaki dominant duyguyu bul
            ilgili_duygular = [d['duygu'] for d in duygu_cizelgesi if baslangic_zamani <= d['zaman'] <= bitis_zamani]
            dominant_duygu = Counter(ilgili_duygular).most_common(1)[0][0] if ilgili_duygular else "BELİRSİZ"

            sonuclar.append({
                'konusmaci': konusmaci_map[blok['speaker_id']],
                'diyalog': cumle_metni,
                'duygu': dominant_duygu,
                'baslangic': baslangic_zamani
            })

    # Sonuçları kronolojik olarak sırala
    sonuclar.sort(key=lambda x: x['baslangic'])
    return sonuclar


def sonucu_word_dosyasina_yaz(dosya_adi, analiz_sonuclari):
    """Analiz sonuçlarını formatlayarak bir Word belgesine yazar."""
    try:
        doc = Document()
        doc.add_heading(f'{VIDEO_DOSYASI} - Duygu Analizli Transkript', level=1)
        doc.add_paragraph(
            f"Bu belge, video dosyasındaki konuşmaların metne dökülmüş halini ve her cümlenin söylendiği andaki dominant yüz ifadesini içerir.")
        doc.add_paragraph()  # Boşluk bırak

        print(f"\n--- Çıktı '{dosya_adi}' Dosyasına Yazılıyor ---")

        for sonuc in analiz_sonuclari:
            line = f"[{sonuc['konusmaci']}] [{sonuc['duygu']}]: {sonuc['diyalog']}"
            print(line)
            doc.add_paragraph(line)

        doc.save(dosya_adi)
        print(f"\nTranskript başarıyla '{dosya_adi}' dosyasına kaydedildi.")
    except Exception as e:
        print(f"\nWord dosyası oluşturulurken bir hata oluştu: {e}")


# --- 3. ANA İŞLEM AKIŞI ---
if __name__ == "__main__":

    # Gerekli kontroller
    if ELEVENLABS_API_KEY == "BURAYA_ELEVENLABS_API_ANAHTARINIZI_YAZIN":
        print("Hata: Lütfen kodun 13. satırına ElevenLabs API anahtarınızı girin.")
        exit()

    if not os.path.exists(VIDEO_DOSYASI):
        print(f"Hata: '{VIDEO_DOSYASI}' video dosyası bulunamadı.")
        exit()

    # spaCy modelini yükle
    try:
        nlp = spacy.load("xx_ent_wiki_sm")
        # DÜZELTME: Cümle sınırlarını tanımak için sentencizer'ı pipeline'a ekle
        nlp.add_pipe('sentencizer')
    except OSError:
        print("\nHata: spaCy modeli 'xx_ent_wiki_sm' bulunamadı.")
        print("Lütfen 'python -m spacy download xx_ent_wiki_sm' komutunu çalıştırarak modeli indirin.")
        exit()

    # 1. Adım: Videodan sesi ayıkla
    if not videodan_sesi_cikar(VIDEO_DOSYASI, GECICI_SES_DOSYASI):
        exit()

    # 2. Adım: Sesi yazıya çevir
    kelime_verisi = sesi_yaziya_cevir(ELEVENLABS_API_KEY, GECICI_SES_DOSYASI)

    # Geçici ses dosyasını bu aşamadan sonra sil
    if os.path.exists(GECICI_SES_DOSYASI):
        os.remove(GECICI_SES_DOSYASI)

    if not kelime_verisi:
        print("ElevenLabs'tan veri alınamadığı için işlem durduruldu.")
        exit()

    # 3. Adım: Videodan duyguları analiz et
    duygu_cizelgesi = videodaki_duygulari_analiz_et(VIDEO_DOSYASI, SANIYEDE_ANALIZ_SAYISI)

    # 4. Adım: Verileri birleştir
    print("\nMetin ve duygu verileri birleştiriliyor...")
    final_output = verileri_birlestir_ve_formatla(kelime_verisi, duygu_cizelgesi, nlp)

    # 5. Adım: Sonucu dosyaya yaz
    if final_output:
        sonucu_word_dosyasina_yaz(CIKTI_DOSYASI, final_output)
    else:
        print("\nSonuç üretilemediği için Word dosyası oluşturulmadı.")

    print("\nTüm işlemler başarıyla tamamlandı.")