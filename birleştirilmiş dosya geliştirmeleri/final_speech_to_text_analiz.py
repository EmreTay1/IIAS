
# -*- coding: utf-8 -*-

# --- GEREKLİ KÜTÜPHANELER ---
import os
import json
import requests
import spacy
from moviepy import VideoFileClip
from docx import Document
from elevenlabs import ElevenLabs
from langchain.text_splitter import RecursiveCharacterTextSplitter

# --- KONFİGÜRASYON ---
# Bu bölümdeki değişkenleri kendi ayarlarınıza göre güncelleyin.

# --- 1. Adım: Ses Deşifre Ayarları ---
ELEVENLABS_API_KEY = "sk_32f4726f5c0d2db9d33a0952645acf8e35986bfad444d311"
INPUT_VIDEO_FILE = "Mülakat Simülasyonu.mp4"
TEMP_AUDIO_FILE = "gecici_ses.wav"
TRANSCRIPT_DOCX_FILE = "mulakat_transkripti.docx"

# --- 2. Adım: Metin Analizi Ayarları ---
LM_STUDIO_API_URL = "http://localhost:1234/v1/chat/completions"
MODEL_NAME = "mistral-7b-instruct-v0.3.Q4_0"
FINAL_ANALYSIS_DOCX_FILE = "analiz_sonucu.docx"
DURATION_THRESHOLD = 20  # Video süresi eşiği (dakika)

# --- YENİ PROMPT'LAR (metin_analizi_chunk.py'den) ---
PROMPT_SCORING_DETAILS = """
Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.

**BÖLÜM 1: PUANLAMA TABLOSU TALİMATLARI**

1.  **Formatlama Kuralı:** Her kriteri, madde imi (`•`) ile başlayan ayrı bir satır olarak ve aşağıdaki formatı **birebir** kullanarak yaz:
    `• Kriter Adı: (Puan/5) - Gerekçe.`

2.  **Gerekçe İçeriği:** Gerekçe bölümünde, puanın nedenini açıklayan, adayın adı ("Can Bey" gibi) ile başlayan kısa ve tanımlayıcı bir cümle kur. Örneğin: `Can Bey'in sorulara verdiği cevapların netliği.`

3.  **Değerlendirilecek Kriterler:**
    * İletişim Becerisi
    * Motivasyon ve Tutku
    * Kültürel Uyum
    * Analitik/Düşünsel Beceriler
    * Profesyonel Tutum
    * Geçmiş Deneyim Uyumu
    * Liderlik ve Girişimcilik
    * Zayıflıklarla Başa Çıkma Yetisi
    * Uzun Vadeli Potansiyel
    * Genel Etki / İzlenim

4.  **Analiz Sonu:** Puanlama listesinin ardından, aşağıdaki iki başlığı ekle:
    * `Genel Ortalama Puan:` Tüm puanların ortalamasını ondalık ayraç olarak virgül kullanarak hesapla. Örnek: `(3,86/5)`
    * `İK Genel Yorum:` Birkaç cümlelik, adayın genel potansiyelini ve ana bulguları özetleyen bir paragraf yaz.
"""

PROMPT_RECRUITER_DETAILS = """
Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.

**BÖLÜM 2: RECRUITER NOTU TALİMATLARI**

1.  -Yapısal Kural: Analizini, aşağıdaki altı başlığın tamamını madde imi (`•`) ile başlayan ayrı satırlar olarak yapılandır.

2.  **İçerik Derinliği:** Her başlığın altına, mülakat özetinden çıkardığın somut bilgilere (projeler, deneyimler, yetenekler) dayanarak detaylı ve profesyonel bir metin yaz. Adayın adını metin içinde uygun yerlerde kullan.

3.  **Başlıklar ve İçerikleri:**
    * `• Aday Adı:` Mülakat metninden adayın adını çıkar.
    * `• Pozisyon:` Mülakat metninden adayın başvurduğu pozisyonu belirle.
    * `• Genel Yorum:` Adayın geçmişi, deneyim süresi ve genel performansı hakkında özet bir paragraf yaz.
    * `• Dikkat Çeken Güçlü Yönler:` Adayın öne çıkan teknik veya sosyal yeteneklerini, projelerden örnekler vererek anlat.
    * `• Geliştirme Alanları:` Adayın hangi konularda kendini geliştirebileceğini ve potansiyel gelişim alanlarını belirt.
    * `• Değerlendirme Önerisi:` Aday için bir sonraki adımları (ikinci görüşme, teknik test vb.) ve gelişimini destekleyecek önerileri (eğitim, kurs vb.) içeren bir paragraf yaz.
"""

# ==============================================================================
# --- YARDIMCI FONKSİYONLAR ---
# ==============================================================================

def konusmacilari_ayir_ve_cumlele(words_data):
    """
    ElevenLabs'ten gelen kelime bazlı diarizasyon verisini işler.
    Konuşmacı etiketlerine göre metni gruplar ve spaCy ile cümlelere ayırır.
    """
    if not words_data:
        print("\nKonuşmacı ayırmak için kelime verisi bulunamadı.")
        return []

    speaker_map = {}
    next_speaker_num = 1
    current_speaker_id_raw = None
    current_speaker_mapped_label = None
    current_speaker_words_buffer = []
    all_speaker_utterances_combined = []

    for word_obj in words_data:
        word_text = word_obj.text
        raw_speaker_id = word_obj.speaker_id

        if raw_speaker_id not in speaker_map:
            speaker_map[raw_speaker_id] = f"Konuşmacı {next_speaker_num}"
            next_speaker_num += 1

        mapped_speaker_label = speaker_map[raw_speaker_id]

        if current_speaker_id_raw is None:
            current_speaker_id_raw = raw_speaker_id
            current_speaker_mapped_label = mapped_speaker_label
        elif raw_speaker_id != current_speaker_id_raw:
            if current_speaker_words_buffer:
                all_speaker_utterances_combined.append({
                    'speaker': current_speaker_mapped_label,
                    'text_combined': "".join(current_speaker_words_buffer).strip()
                })
            current_speaker_id_raw = raw_speaker_id
            current_speaker_mapped_label = mapped_speaker_label
            current_speaker_words_buffer = []

        current_speaker_words_buffer.append(word_text)

    if current_speaker_words_buffer:
        all_speaker_utterances_combined.append({
            'speaker': current_speaker_mapped_label,
            'text_combined': "".join(current_speaker_words_buffer).strip()
        })

    final_diarized_sentences = []
    try:
        nlp = spacy.load("xx_ent_wiki_sm")
        if "sentencizer" not in nlp.pipe_names:
            nlp.add_pipe("sentencizer")
    except OSError:
        print("\nspaCy modeli 'xx_ent_wiki_sm' bulunamadı.")
        print("Lütfen 'python -m spacy download xx_ent_wiki_sm' komutunu çalıştırın.")
        return []
    except Exception as e:
        print(f"spaCy modeli yüklenirken bir hata oluştu: {e}")
        return []

    for entry in all_speaker_utterances_combined:
        doc = nlp(entry['text_combined'])
        sentences_for_speaker = [sent.text.strip() for sent in doc.sents if sent.text.strip()]
        for sent in sentences_for_speaker:
            final_diarized_sentences.append({
                "konusmaci": entry['speaker'],
                "diyalog": sent
            })

    return final_diarized_sentences


def read_text_from_docx(file_path):
    """
    Bir .docx dosyasındaki tüm metni okur ve tek bir metin bloğu olarak döndürür.
    """
    try:
        doc = Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        return '\n'.join(full_text)
    except Exception as e:
        print(f"HATA: '{file_path}' dosyası okunurken bir hata oluştu: {e}")
        return None


def get_llm_analysis(prompt, model_name):
    """
    LM Studio API aracılığıyla yerel LLM'e bir prompt gönderir ve analizi alır.
    """
    payload = {
        "model": model_name,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7,
        "max_tokens": 2048
    }
    try:
        response = requests.post(LM_STUDIO_API_URL, json=payload)
        response.raise_for_status()
        response_data = response.json()
        if "choices" in response_data and response_data["choices"]:
            return response_data["choices"][0]["message"]["content"]
        else:
            print(f"HATA: API yanıtında 'choices' anahtarı bulunamadı veya boş. Yanıt: {response_data}")
            return None
    except requests.exceptions.ConnectionError:
        print("HATA: LM Studio API sunucusuna bağlanılamadı. Lütfen sunucunun çalıştığından emin olun.")
        return None
    except requests.exceptions.HTTPError as http_err:
        print(f"HATA: API isteği sırasında HTTP hatası: {http_err}\nYanıt İçeriği: {response.text}")
        return None
    except Exception as e:
        print(f"HATA: Analiz sırasında beklenmedik bir hata oluştu: {e}")
        return None


def write_analysis_to_docx(file_path, analysis_scoring, analysis_recruiter):
    """
    Analiz sonuçlarını yeni bir .docx dosyasına yazar.
    """
    try:
        doc = Document()
        doc.add_heading('Mülakat Analizi Sonucu', level=0)
        doc.add_heading('1. Aday Değerlendirme Puanlama Tablosu', level=1)
        for line in analysis_scoring.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_paragraph()
        doc.add_heading('2. Recruiter Notu', level=1)
        for line in analysis_recruiter.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.save(file_path)
        print(f"\nAnaliz sonuçları '{file_path}' dosyasına başarıyla kaydedildi.")
    except Exception as e:
        print(f"HATA: Sonuçlar dosyaya yazılırken bir hata oluştu: {e}")


def adim_2_metin_analizi_yap(transcript_file_path):
    """
    ADIM 2: Transkripti DOCX dosyasından okur, LLM'e analiz için gönderir
    ve sonuçları nihai bir DOCX dosyasına kaydeder.
    """
    print(f"--- ADIM 2: '{transcript_file_path}' Dosyasındaki Metin Analiz Ediliyor ---")

    # 2.1. Mülakat transkriptini DOCX dosyasından oku
    interview_text = read_text_from_docx(transcript_file_path)
    if not interview_text or not interview_text.strip():
        print(f"HATA: '{transcript_file_path}' dosyası bulunamadı, bozuk veya boş. Analiz yapılamıyor.")
        return

    print("Mülakat transkripti başarıyla okundu.")

    # 2.2. Puanlama Tablosu için prompt'u hazırla ve LLM'e gönder
    print(f"'{MODEL_NAME}' modeline 'Aday Değerlendirme Puanlama Tablosu' için istek gönderiliyor...")
    prompt_scoring = f"""
    Aşağıdaki mülakat metnini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et. 
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:
    {PROMPT_SCORING_DETAILS}

    --- MÜLAKAT METNİ ---
    {interview_text}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_scoring = get_llm_analysis(prompt_scoring, MODEL_NAME)
    if not analysis_scoring:
        print("Puanlama analizi alınamadı. İşlem durduruluyor.")
        return
    print("Puanlama analizi başarıyla tamamlandı.")

    # 2.3. Recruiter Notu için prompt'u hazırla ve LLM'e gönder
    print(f"'{MODEL_NAME}' modeline 'Recruiter Notu' için istek gönderiliyor...")
    prompt_recruiter = f"""
    Aşağıdaki mülakat metnini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:
    {PROMPT_RECRUITER_DETAILS}

    --- MÜLAKAT METNİ ---
    {interview_text}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_recruiter = get_llm_analysis(prompt_recruiter, MODEL_NAME)
    if not analysis_recruiter:
        print("Recruiter notu analizi alınamadı. İşlem durduruluyor.")
        return
    print("Recruiter notu analizi başarıyla tamamlandı.")

    # 2.4. Birleştirilmiş analiz sonuçlarını tek bir DOCX dosyasına yaz
    write_analysis_to_docx(FINAL_ANALYSIS_DOCX_FILE, analysis_scoring, analysis_recruiter)
    print("--- ADIM 2 TAMAMLANDI ---")


def adim_2_metin_analizi_chunk(transcript_file_path):
    """
    ADIM 2 (Chunk Tabanlı): Transkripti DOCX dosyasından okur, metni parçalara böler,
    her parça için özet oluşturur, özetleri birleştirir ve tek bir nihai analiz yapar.
    """
    print(f"--- ADIM 2 (Chunk Tabanlı): '{transcript_file_path}' Dosyasındaki Metin Analiz Ediliyor ---")

    # 2.1. Mülakat transkriptini DOCX dosyasından oku
    interview_text = read_text_from_docx(transcript_file_path)
    if not interview_text or not interview_text.strip():
        print(f"HATA: '{transcript_file_path}' dosyası bulunamadı, bozuk veya boş. Analiz yapılamıyor.")
        return

    print("Mülakat transkripti başarıyla okundu.")

    # 2.2. Metni parçalara böl
    print("Metin parçalara ayrılıyor...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=4000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(interview_text)
    print(f"Metin {len(chunks)} parçaya bölündü.")

    # 2.3. Her parça için özet oluştur
    chunk_summaries = []
    print("Her bir metin parçası için özetler oluşturuluyor...")
    for i, chunk in enumerate(chunks):
        print(f"Parça {i + 1}/{len(chunks)} işleniyor...")
        prompt_chunk_summary = f"""
        Aşağıdaki mülakat metni parçasını oku. Bu parçadan, adayın aşağıda listelenen yetkinlikleri ile ilgili
        tüm önemli bilgileri, kilit ifadeleri ve somut örnekleri Türkçe olarak maddeler halinde özetle.
        Sadece bu metin parçasında geçen bilgileri kullan.

        Yetkinlikler:
        - İletişim Becerisi
        - Motivasyon ve Tutku
        - Kültürel Uyum
        - Analitik/Düşünsel Beceriler
        - Profesyonel Tutum
        - Geçmiş Deneyim Uyumu
        - Liderlik ve Girişimcilik
        - Zayıflıklarla Başa Çıkma Yetisi
        - Uzun Vadeli Potansiyel
        - Genel Etki / İzlenim

        --- MÜLAKAT METNİ PARÇASI ---
        {chunk}
        --- ÖZETİNİ BURAYA BAŞLAT ---
        """
        summary = get_llm_analysis(prompt_chunk_summary, MODEL_NAME)
        if summary:
            chunk_summaries.append(summary)
        else:
            print(f"Parça {i + 1} için özet alınamadı. İşlem durduruluyor.")
            return

    # 2.4. Özetleri birleştir
    combined_summary = "\n\n---\n\n".join(chunk_summaries)
    print("\nTüm parçaların özetleri başarıyla birleştirildi. Nihai analiz başlıyor.")

    # 2.5. Nihai Puanlama Tablosu analizi
    print("Nihai 'Aday Değerlendirme Puanlama Tablosu' oluşturuluyor...")
    prompt_scoring = f"""
    Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece "1. Aday Değerlendirme Puanlama Tablosu" başlığı altında yapılandır ve detaylandır:
    {PROMPT_SCORING_DETAILS}

    --- MÜLAKAT ÖZETİ METNİ ---
    {combined_summary}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_scoring = get_llm_analysis(prompt_scoring, MODEL_NAME)
    if not analysis_scoring:
        print("Puanlama tablosu analizi oluşturulamadı. İşlem durduruluyor.")
        return

    # 2.6. Nihai Recruiter Notu analizi
    print("Nihai 'Recruiter Notu' oluşturuluyor...")
    prompt_recruiter = f"""
    Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece "2. Recruiter Notu" başlığı altında yapılandır ve detaylandır:
    {PROMPT_RECRUITER_DETAILS}

    --- MÜLAKAT ÖZETİ METNİ ---
    {combined_summary}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_recruiter = get_llm_analysis(prompt_recruiter, MODEL_NAME)
    if not analysis_recruiter:
        print("Recruiter notu analizi oluşturulamadı. İşlem durduruluyor.")
        return

    # 2.7. Birleştirilmiş analiz sonuçlarını DOCX dosyasına yaz
    write_analysis_to_docx(FINAL_ANALYSIS_DOCX_FILE, analysis_scoring, analysis_recruiter)
    print("--- ADIM 2 (Chunk Tabanlı) TAMAMLANDI ---")


# ==============================================================================
# --- ANA AKIŞ FONKSİYONLARI ---
# ==============================================================================

def adim_1_videodan_metne_cevir():
    """
    ADIM 1: Videodan sesi çıkarır, ElevenLabs ile deşifre eder ve sonucu DOCX'e kaydeder.
    Başarılı olursa transkript dosyasının yolunu ve video süresini (dakika cinsinden) döndürür.
    """
    print("--- ADIM 1: Mülakat Videosu Metne Çevriliyor ---")

    # 1.1. Video dosyasının varlığını kontrol et
    if not os.path.exists(INPUT_VIDEO_FILE):
        print(f"HATA: '{INPUT_VIDEO_FILE}' video dosyası bulunamadı. Lütfen kontrol edin.")
        return None, None

    # 1.2. Videodan sesi çıkar ve süreyi hesapla
    try:
        print(f"'{INPUT_VIDEO_FILE}' videosundan ses çıkarılıyor...")
        with VideoFileClip(INPUT_VIDEO_FILE) as video:
            video_duration_minutes = video.duration / 60  # Saniyeyi dakikaya çevir
            video.audio.write_audiofile(TEMP_AUDIO_FILE)
        print(f"Ses başarıyla '{TEMP_AUDIO_FILE}' olarak kaydedildi.")
        print(f"Video süresi: {video_duration_minutes:.2f} dakika")
    except Exception as e:
        print(f"HATA: Video işlenirken bir hata oluştu: {e}")
        return None, None

    # 1.3. ElevenLabs ile sesi yazıya çevir
    words_data = []
    try:
        print(f"'{TEMP_AUDIO_FILE}' dosyası ElevenLabs API'sine gönderiliyor (diarize=True)...")
        elevenlabs_client = ElevenLabs(api_key=ELEVENLABS_API_KEY)
        with open(TEMP_AUDIO_FILE, "rb") as audio_file:
            response = elevenlabs_client.speech_to_text.convert(
                file=audio_file,
                model_id="scribe_v1",
                diarize=True,
            )
            words_data = response.words
        print("Ses, ElevenLabs tarafından başarıyla deşifre edildi.")
    except Exception as e:
        print(f"HATA: ElevenLabs STT işlemi sırasında hata: {e}")
        return None, None
    finally:
        # 1.4. Geçici ses dosyasını sil
        if os.path.exists(TEMP_AUDIO_FILE):
            os.remove(TEMP_AUDIO_FILE)
            print(f"Geçici ses dosyası '{TEMP_AUDIO_FILE}' silindi.")

    # 1.5. Metni konuşmacılara göre ayır
    if not words_data:
        print("HATA: ElevenLabs'tan kelime bazlı veri alınamadı. İşlem durduruluyor.")
        return None, None

    diarized_output = konusmacilari_ayir_ve_cumlele(words_data)
    if not diarized_output:
        print("HATA: Konuşmacı ayrımı yapılamadı veya boş sonuç döndü.")
        return None, None

    # 1.6. Sonucu Word dosyasına yaz
    try:
        doc = Document()
        doc.add_heading('Mülakat Transkripti', level=1)
        for entry in diarized_output:
            line = f"[{entry['konusmaci']}]: {entry['diyalog']}"
            doc.add_paragraph(line)
        doc.save(TRANSCRIPT_DOCX_FILE)
        print(f"Transkript başarıyla '{TRANSCRIPT_DOCX_FILE}' dosyasına kaydedildi.")
        print("--- ADIM 1 TAMAMLANDI ---\n")
        return TRANSCRIPT_DOCX_FILE, video_duration_minutes
    except Exception as e:
        print(f"HATA: Transkript Word dosyasına yazılırken hata oluştu: {e}")
        return None, None


# ==============================================================================
# --- BETİĞİ ÇALIŞTIR ---
# ==============================================================================

if __name__ == "__main__":
    print("===== MÜLAKAT VİDEOSU ANALİZ SÜRECİ BAŞLATILDI =====")

    # Adım 1'i çalıştır: Videoyu metne çevir ve süreyi al
    transcript_file, video_duration = adim_1_videodan_metne_cevir()

    # Adım 1 başarılı olduysa, Adım 2'ye geç
    if transcript_file and video_duration is not None:
        # Video süresine göre analiz yöntemini seç
        if video_duration > DURATION_THRESHOLD:
            print(f"Video süresi {video_duration:.2f} dakika, chunk tabanlı analiz kullanılıyor.")
            adim_2_metin_analizi_chunk(transcript_file)
        else:
            print(f"Video süresi {video_duration:.2f} dakika, standart analiz kullanılıyor.")
            adim_2_metin_analizi_yap(transcript_file)

        # (İsteğe bağlı) Ara transkript dosyasını sil
        # os.remove(transcript_file)
        # print(f"\nAra transkript dosyası '{transcript_file}' silindi.")

    else:
        print("\nSüreç, Adım 1'deki bir hata nedeniyle durduruldu. Analiz yapılamadı.")

    print("\n===== SÜREÇ TAMAMLANDI =====")